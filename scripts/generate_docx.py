#!/usr/bin/env python3
"""
Genera agent3_article.docx con estilos de titulo y parrafo formateados
para Manager in Motion — SEO Pipeline 2026-06-08
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

OUTPUT_PATH = "/home/user/mimo-seo/output/agent3_article.docx"

# ── Paleta de marca Manager in Motion ──────────────────────────────────────
COLOR_TITULO    = RGBColor(0x1A, 0x1A, 0x2E)   # azul marino oscuro
COLOR_H2        = RGBColor(0x16, 0x21, 0x3E)   # azul profundo
COLOR_H3        = RGBColor(0x0F, 0x3E, 0x65)   # azul medio
COLOR_CUERPO    = RGBColor(0x2C, 0x2C, 0x2C)   # gris carbón
COLOR_META      = RGBColor(0x55, 0x55, 0x55)   # gris medio
COLOR_CTA_BG    = RGBColor(0xE8, 0xF4, 0xFD)   # azul claro fondo
COLOR_ACENTO    = RGBColor(0x00, 0x7B, 0xC2)   # azul Manager in Motion
COLOR_FUENTES   = RGBColor(0x77, 0x77, 0x77)   # gris claro fuentes

FONT_SANS = "Calibri"
FONT_SERIF = "Georgia"


def set_cell_background(cell, hex_color):
    """Añade color de fondo a una celda de tabla."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def add_heading(doc, text, level, color):
    """Añade encabezado con color y fuente controlados."""
    p = doc.add_paragraph()
    p.style = f'Heading {level}'
    run = p.add_run(text)
    run.font.color.rgb = color
    run.font.name = FONT_SANS
    run.font.bold = True
    if level == 1:
        run.font.size = Pt(22)
    elif level == 2:
        run.font.size = Pt(16)
    elif level == 3:
        run.font.size = Pt(13)
    p.paragraph_format.space_before = Pt(18 if level <= 2 else 12)
    p.paragraph_format.space_after = Pt(6)
    return p


def add_body(doc, text, bold=False, italic=False, size=11, color=None):
    """Añade párrafo de cuerpo con formato."""
    p = doc.add_paragraph()
    p.style = 'Normal'
    run = p.add_run(text)
    run.font.name = FONT_SANS
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = color or COLOR_CUERPO
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = Pt(16)
    p.paragraph_format.left_indent = Inches(0)
    return p


def add_meta_block(doc, label, value):
    """Añade bloque de metadato con etiqueta y valor."""
    p = doc.add_paragraph()
    run_label = p.add_run(f"{label}: ")
    run_label.font.name = FONT_SANS
    run_label.font.size = Pt(10)
    run_label.font.bold = True
    run_label.font.color.rgb = COLOR_ACENTO
    run_val = p.add_run(value)
    run_val.font.name = FONT_SANS
    run_val.font.size = Pt(10)
    run_val.font.color.rgb = COLOR_META
    p.paragraph_format.space_after = Pt(3)
    return p


def add_cta_box(doc, text):
    """Añade caja de CTA destacada con fondo azul claro."""
    table = doc.add_table(rows=1, cols=1)
    table.style = 'Table Grid'
    cell = table.cell(0, 0)
    set_cell_background(cell, 'D6EAF8')
    cell.width = Inches(6)
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.left_indent = Inches(0.2)
    run = p.add_run(text)
    run.font.name = FONT_SANS
    run.font.size = Pt(11)
    run.font.bold = True
    run.font.color.rgb = COLOR_H2
    doc.add_paragraph()


def add_separator(doc):
    """Añade línea divisoria horizontal."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '007BC2')
    pBdr.append(bottom)
    pPr.append(pBdr)


def add_faq_item(doc, question, answer):
    """Añade item FAQ con pregunta en negrita azul y respuesta normal."""
    p = doc.add_paragraph()
    run_q = p.add_run(question)
    run_q.font.name = FONT_SANS
    run_q.font.size = Pt(11)
    run_q.font.bold = True
    run_q.font.color.rgb = COLOR_ACENTO
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)

    p2 = doc.add_paragraph()
    run_a = p2.add_run(answer)
    run_a.font.name = FONT_SANS
    run_a.font.size = Pt(11)
    run_a.font.color.rgb = COLOR_CUERPO
    p2.paragraph_format.space_after = Pt(6)
    p2.paragraph_format.left_indent = Inches(0.2)


def add_source_item(doc, text):
    """Añade elemento de fuente en itálica gris."""
    p = doc.add_paragraph()
    run = p.add_run(f"• {text}")
    run.font.name = FONT_SANS
    run.font.size = Pt(9)
    run.font.italic = True
    run.font.color.rgb = COLOR_FUENTES
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Inches(0.2)


# ── CONSTRUCCIÓN DEL DOCUMENTO ──────────────────────────────────────────────

doc = Document()

# Márgenes de página
for section in doc.sections:
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)

# ── CABECERA DE METADATOS ───────────────────────────────────────────────────
p_brand = doc.add_paragraph()
run_brand = p_brand.add_run("MANAGER IN MOTION  |  Artículo SEO  |  2026-06-08")
run_brand.font.name = FONT_SANS
run_brand.font.size = Pt(9)
run_brand.font.bold = True
run_brand.font.color.rgb = COLOR_ACENTO
p_brand.paragraph_format.space_after = Pt(4)

add_meta_block(doc, "Título SEO", "Directivo interino para transformación digital: la solución al déficit de talento IA en España")
add_meta_block(doc, "Meta descripción", "El 67% de empresas españolas no tiene el talento para ejecutar su hoja de ruta IA. Descubre cómo un directivo interino especializado cierra esa brecha en semanas, no en meses.")
add_meta_block(doc, "Slug URL", "directivo-interino-transformacion-digital-ia-espana")
add_meta_block(doc, "Keyword principal", "directivo interino transformación digital")

add_separator(doc)
doc.add_paragraph()

# ── H1 ──────────────────────────────────────────────────────────────────────
add_heading(doc, "Por qué el directivo interino es la respuesta al déficit de liderazgo en transformación IA que frena al 67% de las empresas españolas", 1, COLOR_TITULO)

doc.add_paragraph()

# ── INTRODUCCIÓN ─────────────────────────────────────────────────────────────
add_body(doc, "El problema no es la falta de inversión. El problema es la falta del directivo que ejecute.")

add_body(doc, "Según el KPMG Global Tech Report 2026 —edición España, publicado en abril de 2026 sobre más de 100 respuestas de directivos españoles—, el 67% de las organizaciones en España reconoce que carece del talento necesario para ejecutar su estrategia de transformación digital. Seis de cada diez aspiran a alcanzar la máxima madurez en inteligencia artificial en 2026. Solo el 23% lo ha conseguido. El 81% ya invierte en agentes de IA. La brecha no está entre querer y pagar: está entre pagar y poder ejecutar.")

add_body(doc, "Este artículo no explica qué es la inteligencia artificial. Se dirige al CEO, CTO o CHRO que ya tiene aprobado un presupuesto, una hoja de ruta o un mandato de transformación y se enfrenta a la misma pregunta semana tras semana: ¿quién lo lidera?")

add_body(doc, "La respuesta más eficiente disponible hoy en el mercado español no es contratar un CTO permanente. Es contratar un directivo interino especializado en transformación digital.")

add_separator(doc)

# ── H2-1 ─────────────────────────────────────────────────────────────────────
add_heading(doc, "El dato que define el reto: el 67% de las empresas españolas no puede ejecutar su transformación IA", 2, COLOR_H2)

add_heading(doc, "Lo que dice el KPMG Global Tech Report 2026 sobre el talento tecnológico en España", 3, COLOR_H3)

add_body(doc, "El informe de KPMG es explícito: la escasez de talento es más crítica en España que en el promedio global. Mientras a nivel mundial el porcentaje de organizaciones que admite no tener el talento necesario se sitúa en el 53%, en España asciende al 67%. No es una diferencia marginal. Refleja un mercado laboral tecnológico estructuralmente insuficiente para la velocidad a la que las organizaciones quieren moverse.")

add_body(doc, "El dato no es anecdótico. Está construido sobre respuestas de directivos de tecnología de 27 países, y la edición española recoge la realidad de más de cien organizaciones de ocho sectores distintos. El 90% de esas organizaciones considera que la gestión de agentes de IA será una habilidad crítica en cinco años. El 91% planea ampliar alianzas tecnológicas en los próximos doce meses. La dirección está clara; el talento ejecutivo para transitarla, no.")

add_heading(doc, "La brecha entre ambición digital y capacidad de ejecución: por qué el problema no es la inversión, sino el liderazgo", 3, COLOR_H3)

add_body(doc, "El diagnóstico de KPMG coincide con lo que los CODIR de empresas medianas españolas viven en primera persona: la inversión tecnológica está aprobada, los proyectos están en marcha, pero el liderazgo ejecutivo que los conduce es insuficiente, inexperto o inexistente. Los equipos de tecnología tienen capacidad técnica; lo que no tienen es un directivo con experiencia en transformaciones a escala que sepa traducir la agenda IA en decisiones organizativas, priorizar iniciativas, gestionar resistencias y rendir cuentas ante el CEO.")

add_body(doc, "Esa figura —el directivo que ha pilotado antes una transformación de esta complejidad— no se fabrica en seis meses ni se improvisa con una promoción interna. Se contrata. Y la pregunta real es: ¿con qué modelo?")

add_separator(doc)

# ── H2-2 ─────────────────────────────────────────────────────────────────────
add_heading(doc, "Por qué la contratación permanente no resuelve este problema (al menos no a tiempo)", 2, COLOR_H2)

add_heading(doc, "El coste real de un proceso de selección de CTO o CDO: tiempo, recursos y riesgo de error", 3, COLOR_H3)

add_body(doc, "Un proceso de selección ejecutiva para un CTO o CDO de nivel en España tiene una duración real de cuatro a seis meses desde el inicio del briefing hasta la incorporación efectiva. Ese plazo incluye la definición del perfil, la búsqueda de candidatos, las rondas de entrevistas, la evaluación de referencias, la negociación de condiciones y el preaviso del candidato seleccionado. En proyectos de transformación IA con plazos comprometidos ante inversores o consejos de administración, seis meses no son una opción.")

add_body(doc, "A eso hay que añadir el coste de equivocarse. Una contratación permanente equivocada en un rol C-suite —con los costes de liquidación, la pérdida de momentum en el proyecto y el coste reputacional de empezar de nuevo— puede costar entre dos y cuatro años de salario del directivo en cuestión. En perfiles con salarios de 150.000 a 250.000 euros anuales, el impacto es significativo.")

add_heading(doc, "La incertidumbre macroeconómica en España como freno a compromisos directivos permanentes", 3, COLOR_H3)

add_body(doc, "El entorno económico de 2026 añade otro factor de peso. Según PwC España en sus Claves Económicas de junio de 2026, el BCE afronta una posible subida de tipos ante una inflación que en España supera el 3%, y el apetito inversor se modera. El 55% de los empresarios y directivos españoles no espera cambios significativos en su entorno económico en los próximos doce meses, según KPMG Perspectivas 2026. En ese contexto, comprometerse con un salario directivo de alto nivel, con todos sus costes asociados, es una decisión que muchos CODIR posponen indefinidamente, dejando el hueco de liderazgo tecnológico sin cubrir.")

add_heading(doc, "El modelo interim-to-permanent: probar antes de contratar de forma definitiva", 3, COLOR_H3)

add_body(doc, "Una vía intermedia que gana terreno en los mercados europeos maduros es el modelo interim-to-permanent: el directivo interino se incorpora en modo misión y, si la valoración es positiva tras un período acordado, se formaliza su contratación permanente. En Francia, cerca del 10% de las misiones interinas concluyen de esta forma. Heidrick & Struggles proyecta que este modelo crecerá un 30-35% en el mercado de Reino Unido en 2026. El modelo elimina el riesgo más costoso de la contratación ejecutiva: contratar definitivamente a alguien que no has visto actuar en tu contexto real.")

add_separator(doc)

# ── H2-3 ─────────────────────────────────────────────────────────────────────
add_heading(doc, "Qué hace exactamente un directivo interino de transformación digital o IA", 2, COLOR_H2)

add_heading(doc, "El CTO interino: estabilización tecnológica, arquitectura de datos y liderazgo de equipos técnicos", 3, COLOR_H3)

add_body(doc, "Un CTO interino asume la dirección del área tecnológica con plena responsabilidad ejecutiva. No asesora: decide. Dirige los equipos de ingeniería y datos, define o revisa la arquitectura tecnológica, evalúa la deuda técnica acumulada, establece prioridades en el roadmap y representa la función tecnológica ante el CEO y el consejo. En misiones de transformación IA, su foco es doble: asegurar que la infraestructura tecnológica soporta los casos de uso priorizados, y garantizar que los equipos internos están alineados y capacitados para sostener los avances después de que él salga.")

add_body(doc, "La diferencia con un consultor tecnológico es fundamental: el consultor entrega un informe y se va. El CTO interino ejecuta y responde de los resultados.")

add_heading(doc, "El CDO interino: estrategia de datos, gobernanza IA y gestión del cambio cultural", 3, COLOR_H3)

add_body(doc, "El Chief Data Officer interino trabaja en la dimensión que más organizaciones subestiman: los datos como activo y la cultura como condición de éxito. Sin una estrategia de datos coherente, ningún proyecto de IA escala. Sin gestión del cambio cultural, ningún proyecto de IA se adopta. El CDO interino construye el gobierno del dato, define los criterios de calidad y seguridad de los modelos de IA, establece marcos de responsabilidad para el uso de IA generativa y trabaja con los mandos intermedios para reducir la resistencia al cambio que KPMG identifica como uno de los principales frenos a la madurez IA.")

add_heading(doc, "El Director de Transformación interino: integración de iniciativas, priorización de roadmap y entrega de resultados", 3, COLOR_H3)

add_body(doc, "Cuando la empresa tiene múltiples iniciativas de transformación digital en marcha —algunas dentro de tecnología, otras en operaciones, marketing o finanzas— pero carece de una figura que las integre y las lleve a buen puerto, el Director de Transformación interino es el perfil adecuado. Su misión es horizontal: garantiza que los proyectos no funcionan en silos, que las prioridades están alineadas con la estrategia de negocio, que hay un criterio único de medición del avance y que alguien rinde cuentas de la transformación como un todo ante la dirección general.")

add_separator(doc)

# ── H2-4 ─────────────────────────────────────────────────────────────────────
add_heading(doc, "La evidencia europea: los mercados maduros ya lo han adoptado como estándar", 2, COLOR_H2)

add_heading(doc, "Reino Unido: +58% en nombramientos de liderazgo interino digital desde 2023", 3, COLOR_H3)

add_body(doc, "El mercado más desarrollado de Europa para el interim management es también el que ofrece los datos más reveladores sobre la dirección del mercado. Según Heidrick & Struggles en su informe de 2026 sobre talento interino en empresas del Reino Unido, los nombramientos de liderazgo interino en las funciones de digital, datos y tecnología han crecido un 58% desde 2023. El 89% de los directivos interinos activos ya emplea herramientas de IA en sus propias misiones, comprimiendo el tiempo entre su incorporación y la generación de impacto visible. Una cuarta parte de todas las solicitudes de directivos interinos, independientemente de la función, incluyen ahora un componente de digital, datos o IA.")

add_heading(doc, "Francia: el directivo interino como agente de cambio cultural en la adopción de IA", 3, COLOR_H3)

add_body(doc, "Francia es el mercado de referencia en Europa continental para el interim management. Aunque 2025 registró un ajuste de volumen del -15%, la percepción de valor no se erosionó: las tarifas para perfiles de alta complejidad se mantienen por encima de los 1.300 euros diarios. Lo más relevante es el cambio de rol: los directivos interinos han pasado de ser gestores de crisis a convertirse en agentes de cambio cultural IA, con la misión de normalizar el aprendizaje continuo y vencer la resistencia de la capa media directiva ante la transformación.")

add_heading(doc, "Alemania: ciberseguridad y automatización industrial como vectores de demanda interina", 3, COLOR_H3)

add_body(doc, "El DDIM, la asociación alemana de interim management, publicó su Marktstudie 2026 con datos que confirman la transformación sectorial del mercado. El sector del automóvil ha caído del primer al cuarto puesto en demanda de directivos interinos. Crecen con fuerza la energía, la infraestructura y la manufactura general. Las misiones de ciberseguridad y automatización industrial son las de mayor crecimiento relativo. El mercado alemán, con 2.700 millones de euros de volumen y 12.500 interinos activos, funciona como espejo de lo que el mercado español vivirá en los próximos tres a cinco años.")

add_heading(doc, "España en el horizonte de 12-18 meses: qué podemos anticipar del patrón europeo", 3, COLOR_H3)

add_body(doc, "El patrón europeo es consistente: los mercados que han alcanzado madurez en interim management lo usan de forma preferente para misiones de transformación y cambio, no solo para cubrir vacantes urgentes. España lleva un retraso estructural de tres a cinco años respecto a Francia y Alemania, y de cinco a siete respecto al Reino Unido. La presión de la transformación IA es simultánea en todos los mercados, y las empresas españolas que no encuentren el liderazgo ejecutivo internamente recurrirán al modelo interino, exactamente como ya hacen sus homólogas europeas.")

add_separator(doc)

# ── H2-5 ─────────────────────────────────────────────────────────────────────
add_heading(doc, "Cuándo contratar un directivo interino de transformación digital — cinco situaciones concretas", 2, COLOR_H2)

add_body(doc, "Estas cinco situaciones no son teóricas. Son los mandatos reales que se abren en el mercado hoy.")

add_heading(doc, "Cuando tienes aprobado un proyecto IA pero no el equipo directivo para ejecutarlo", 3, COLOR_H3)
add_body(doc, "El presupuesto está aprobado, el proveedor tecnológico está contratado, los casos de uso están definidos. Pero el proyecto no avanza porque no hay nadie con autoridad ejecutiva real para tomarlo y llevarlo. Los equipos técnicos hacen lo que pueden; los mandos intermedios esperan instrucciones claras. Un directivo interino de transformación digital toma el liderazgo del proyecto en las primeras semanas y lo lleva a buen puerto.")

add_heading(doc, "Cuando necesitas acelerar una hoja de ruta digital antes de una operación corporativa o ronda de inversión", 3, COLOR_H3)
add_body(doc, "Los inversores y los fondos de private equity valoran la madurez digital en sus due diligence. Una empresa que demuestra avance real en IA y datos negocia desde una posición más sólida. El directivo interino permite comprimir el tiempo entre la situación actual y la que el mercado valora, sin el riesgo de una contratación permanente que puede no encajar tras el cierre de la operación.")

add_heading(doc, "Cuando tu CTO actual no tiene el perfil para liderar la transición IA y necesitas un puente ejecutivo", 3, COLOR_H3)
add_body(doc, "Muchos CTOs de empresas medianas tienen un perfil de gestión de sistemas e infraestructura que no cubre la dimensión estratégica de la transformación IA: gobernanza de datos, arquitectura de modelos, gestión del cambio cultural, comunicación con el consejo. El directivo interino actúa como puente: cubre la brecha mientras el CTO se desarrolla, o mientras se lleva a cabo la búsqueda del sucesor adecuado.")

add_heading(doc, "Cuando el proceso de selección permanente se extiende más allá de lo que el negocio puede esperar", 3, COLOR_H3)
add_body(doc, "El proceso de búsqueda lleva meses, el mercado de perfiles senior en IA está tensionado y el proyecto no puede esperar. El directivo interino se incorpora en dos a cuatro semanas, estabiliza la situación y garantiza que el proyecto no pierde momentum mientras la selección permanente avanza en paralelo.")

add_heading(doc, "Cuando buscas reducir el riesgo de una decisión directiva de alto coste e impacto irreversible", 3, COLOR_H3)
add_body(doc, "Una contratación ejecutiva equivocada en transformación digital puede costar años de retraso y recursos desperdiciados. El modelo interino permite validar el perfil, el enfoque y la compatibilidad cultural antes de tomar una decisión permanente de alto riesgo. El directivo interino no es la solución definitiva en todos los casos: a veces es el paso que hace posible que la solución definitiva sea la correcta.")

doc.add_paragraph()
add_cta_box(doc, "¿Tu empresa tiene una hoja de ruta IA pero no el directivo que la ejecute?\nEn Manager in Motion identificamos el perfil exacto que necesitas en menos de dos semanas.\nHabla con nosotros sin compromiso: www.managerinmotion.eu/contacto")

add_separator(doc)

# ── H2-6 ─────────────────────────────────────────────────────────────────────
add_heading(doc, "El proceso en la práctica — cómo funciona una misión de transformación digital con directivo interino", 2, COLOR_H2)

add_heading(doc, "De la identificación del perfil a la incorporación: plazos reales en el mercado español", 3, COLOR_H3)
add_body(doc, "El plazo entre el primer contacto con una firma de gestión interina ejecutiva y la incorporación efectiva del directivo oscila entre dos y cuatro semanas en el mercado español. Ese plazo incluye la definición del perfil con la empresa cliente, la identificación de los candidatos más adecuados de la red de directivos disponibles, las entrevistas de selección y la firma del acuerdo de misión. Comparado con los cuatro a seis meses de un proceso de selección permanente, la diferencia es operativamente decisiva.")

add_heading(doc, "Los primeros 30 días: diagnóstico, alineación y primeras acciones con impacto visible", 3, COLOR_H3)
add_body(doc, "El estándar de calidad en los primeros 30 días de una misión interina es claro: diagnóstico del estado real de la transformación, identificación de los cuellos de botella críticos, alineación con el CEO y el CODIR sobre prioridades y métricas de éxito, y al menos una acción visible con impacto concreto en el proyecto. La velocidad de incorporación no es una promesa comercial: es una competencia central del directivo interino experimentado, que ha pasado por suficientes procesos de onboarding en contextos complejos como para hacerlo sistemáticamente bien.")

add_heading(doc, "Medición de resultados y criterios de éxito en una misión de transformación IA", 3, COLOR_H3)
add_body(doc, "Los criterios de éxito se definen antes del inicio del mandato: hitos de proyecto, KPIs de adopción tecnológica, grado de avance del roadmap, estado de la gobernanza de datos, nivel de capacitación del equipo interno y, cuando procede, condiciones para la entrega de la función a un sucesor permanente. La medición es parte del modelo: el directivo interino no trabaja con objetivos difusos.")

add_separator(doc)

# ── H2-7 FAQ ─────────────────────────────────────────────────────────────────
add_heading(doc, "Preguntas frecuentes sobre el directivo interino de transformación digital", 2, COLOR_H2)

add_faq_item(doc,
    "¿Qué diferencia hay entre un CTO interino y un consultor de transformación digital?",
    "El CTO interino asume responsabilidad ejecutiva real: reporta al CEO, dirige equipos, toma decisiones y responde de los resultados. El consultor recomienda sin ejecutar. En una transformación IA, la diferencia es determinante: el directivo interino lleva la iniciativa, el consultor entrega un informe. Cuando el proyecto necesita liderazgo —no solo diagnóstico— la figura adecuada es el directivo interino."
)

add_faq_item(doc,
    "¿En cuánto tiempo puede incorporarse un directivo interino de transformación digital a mi empresa?",
    "En el mercado español, los plazos de incorporación rondan las dos a cuatro semanas desde que se define el perfil. Frente a los cuatro a seis meses de un proceso de selección permanente, la ventaja temporal es decisiva cuando la hoja de ruta IA tiene plazos comprometidos ante inversores, consejo o la propia dirección general."
)

add_faq_item(doc,
    "¿Tiene sentido contratar un CDO o CTO interino para una empresa mediana de menos de 100 millones de facturación?",
    "Sí, y es precisamente el segmento donde mayor impacto relativo genera. Las empresas medianas son las que más sufren la brecha de talento tecnológico: no tienen el perfil internamente y el coste de una contratación permanente de nivel C-suite es difícil de justificar sin certeza sobre el ROI. El directivo interino resuelve exactamente esa ecuación: acceso a experiencia de primer nivel, coste proporcional a la duración del mandato, sin compromisos permanentes."
)

add_faq_item(doc,
    "¿Cómo se mide el éxito de una misión de transformación digital con un directivo interino?",
    "Los criterios se acuerdan antes del inicio del mandato: hitos de proyecto, KPIs de adopción tecnológica, grado de avance del roadmap, capacitación del equipo y, cuando procede, entrega de la función a un sucesor permanente. No hay ambigüedad sobre lo que se espera ni sobre cómo se evalúa."
)

add_faq_item(doc,
    "¿Puede un directivo interino de transformación digital liderar la implementación de IA generativa en mi empresa?",
    "Sí. Los mandatos vinculados a IA generativa son los de mayor crecimiento en Europa en 2026. Los directivos interinos especializados trabajan tanto en la dimensión tecnológica —selección de herramientas, arquitectura, seguridad— como en la organizativa: gestión del cambio, formación de equipos, gobernanza del dato y marcos de uso responsable."
)

add_faq_item(doc,
    "¿Cuál es el coste de un directivo interino de transformación digital comparado con una contratación permanente?",
    "La tarifa diaria de un directivo interino de alto perfil tecnológico en España oscila entre 900 y 1.500 euros, dependiendo del nivel del perfil y la complejidad del mandato. Para una misión de seis meses a cuatro días semanales, el coste total resulta considerablemente inferior al salario anual más bonus más costes de selección y Seguridad Social de un CTO o CDO permanente. Y sin el coste del riesgo de una contratación equivocada."
)

add_faq_item(doc,
    "¿El directivo interino puede convertirse en contratación permanente si la misión va bien?",
    "Sí. El modelo interim-to-permanent, que ya representa cerca del 10% de las misiones en mercados maduros como Francia, y que Heidrick & Struggles proyecta que crecerá un 30-35% en el mercado de Reino Unido en 2026, permite evaluar al directivo en contexto real antes de formalizar una contratación permanente. Es la forma más eficiente de eliminar el riesgo de incorporación equivocada en un rol crítico."
)

doc.add_paragraph()

# ── CTA FINAL ────────────────────────────────────────────────────────────────
add_separator(doc)
doc.add_paragraph()
add_cta_box(doc, "¿Tu empresa tiene una hoja de ruta IA pero no el directivo que la ejecute?\nHabla con nosotros. En Manager in Motion identificamos el perfil exacto que necesitas en menos de dos semanas.\nConsulta sin compromiso: www.managerinmotion.eu/contacto")

doc.add_paragraph()

# ── FUENTES ──────────────────────────────────────────────────────────────────
add_separator(doc)
p_fuentes_title = doc.add_paragraph()
run_ft = p_fuentes_title.add_run("Fuentes")
run_ft.font.name = FONT_SANS
run_ft.font.size = Pt(10)
run_ft.font.bold = True
run_ft.font.color.rgb = COLOR_META
p_fuentes_title.paragraph_format.space_before = Pt(10)
p_fuentes_title.paragraph_format.space_after = Pt(4)

sources = [
    "KPMG España — Global Tech Report 2026, edición España (abril 2026). kpmg.com/es",
    "Heidrick & Struggles — How interim talent is delivering transformation for UK businesses in 2026. heidrick.com",
    "DDIM — Marktstudie 2026: Interim Management behauptet sich in anspruchsvollem Marktumfeld. ddim.de",
    "PwC España — Claves Económicas, junio 2026. pwc.es",
    "KPMG España — Perspectivas 2026: Previsiones económicas y empresariales (marzo 2026). kpmg.com/es",
]
for s in sources:
    add_source_item(doc, s)

# ── GUARDAR ──────────────────────────────────────────────────────────────────
os.makedirs(os.path.dirname(OUTPUT_PATH), exist_ok=True)
doc.save(OUTPUT_PATH)
print(f"Documento generado: {OUTPUT_PATH}")
