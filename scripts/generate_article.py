#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Generador del artículo Agent 3 — Manager in Motion
Produce output/agent3_article.docx con estilos Heading 1/2/3 y Normal
"""

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

OUTPUT_PATH = "/home/user/mimo-seo/output/agent3_article.docx"

# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def add_horizontal_rule(doc):
    """Adds a thin horizontal line paragraph."""
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'CCCCCC')
    pBdr.append(bottom)
    pPr.append(pBdr)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)


def add_cta_box(doc, title_text, body_text, button_text):
    """Adds a styled CTA paragraph with a left border."""
    # Title line
    p_title = doc.add_paragraph()
    p_title.paragraph_format.left_indent = Cm(0.8)
    p_title.paragraph_format.space_before = Pt(10)
    p_title.paragraph_format.space_after = Pt(2)
    run = p_title.add_run(title_text)
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x1A, 0x56, 0x9E)  # corporate blue
    _add_left_border(p_title, color='1A569E')

    # Body line
    p_body = doc.add_paragraph()
    p_body.paragraph_format.left_indent = Cm(0.8)
    p_body.paragraph_format.space_before = Pt(0)
    p_body.paragraph_format.space_after = Pt(2)
    r_body = p_body.add_run(body_text)
    r_body.font.size = Pt(10)
    _add_left_border(p_body, color='1A569E')

    # Button simulation
    p_btn = doc.add_paragraph()
    p_btn.paragraph_format.left_indent = Cm(0.8)
    p_btn.paragraph_format.space_before = Pt(4)
    p_btn.paragraph_format.space_after = Pt(10)
    r_btn = p_btn.add_run(f"[ {button_text} ]")
    r_btn.bold = True
    r_btn.font.size = Pt(10)
    r_btn.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    # Background shading for button effect not easily done inline; use bold blue text
    r_btn.font.color.rgb = RGBColor(0x1A, 0x56, 0x9E)
    _add_left_border(p_btn, color='1A569E')


def _add_left_border(paragraph, color='1A569E', sz='18'):
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    left = OxmlElement('w:left')
    left.set(qn('w:val'), 'single')
    left.set(qn('w:sz'), sz)
    left.set(qn('w:space'), '12')
    left.set(qn('w:color'), color)
    pBdr.append(left)
    pPr.append(pBdr)


def set_body_style(paragraph, size_pt=11, space_before=4, space_after=8):
    paragraph.paragraph_format.space_before = Pt(space_before)
    paragraph.paragraph_format.space_after = Pt(space_after)
    for run in paragraph.runs:
        run.font.size = Pt(size_pt)
        run.font.name = 'Calibri'


# ---------------------------------------------------------------------------
# Document construction
# ---------------------------------------------------------------------------

doc = Document()

# --- Page margins ---
for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3)
    section.right_margin = Cm(3)

# ---------------------------------------------------------------------------
# SECCIÓN 0 — METADATOS SEO
# ---------------------------------------------------------------------------
p = doc.add_paragraph()
p.style = doc.styles['Heading 1']
p.clear()
run = p.add_run('METADATOS SEO')
run.font.size = Pt(13)
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

meta_lines = [
    ('Título SEO', 'Directivo interino en la empresa familiar: guía para el relevo'),
    ('Meta descripción', 'Descubre cómo un directivo interino puede liderar el relevo generacional en tu empresa familiar. Datos, claves prácticas y cómo activarlo. Consulta sin compromiso.'),
    ('Slug URL', 'directivo-interino-empresa-familiar-relevo-generacional'),
    ('H1', 'Directivo interino en la empresa familiar: la solución que transforma el relevo generacional en una ventaja competitiva'),
    ('Keyword principal', 'directivo interino empresa familiar'),
    ('Keywords secundarias', 'gestión interina relevo generacional | manager de transición sucesión empresarial | directivo interino pyme española | interim management empresa familiar España'),
]
for label, value in meta_lines:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    run_label = p.add_run(f'{label}: ')
    run_label.bold = True
    run_label.font.size = Pt(10)
    run_label.font.name = 'Calibri'
    run_value = p.add_run(value)
    run_value.font.size = Pt(10)
    run_value.font.name = 'Calibri'

add_horizontal_rule(doc)
doc.add_paragraph()  # spacer

# ---------------------------------------------------------------------------
# H1
# ---------------------------------------------------------------------------
h1 = doc.add_heading('Directivo interino en la empresa familiar: la solución que transforma el relevo generacional en una ventaja competitiva', level=1)
h1.paragraph_format.space_before = Pt(12)
h1.paragraph_format.space_after = Pt(10)

# ---------------------------------------------------------------------------
# INTRODUCCIÓN
# ---------------------------------------------------------------------------
intro_text = (
    "En España, más de 1,1 millones de empresas familiares deberán afrontar una "
    "decisión crítica en los próximos años: ¿quién toma el relevo? Según el Instituto "
    "de la Empresa Familiar, el 89 % del tejido empresarial español tiene estructura "
    "familiar y genera el 67 % del empleo privado. Un tercio de sus propietarios "
    "alcanzará la edad de jubilación en esta misma década."
)
p = doc.add_paragraph(intro_text)
set_body_style(p)

intro_text2 = (
    "La mayoría no tiene un plan. Algunas tienen un sucesor sobre el papel, pero no "
    "en la práctica. Otras esperan a que la situación se resuelva sola —lo que "
    "raramente ocurre. El instinto natural es llamar al abogado de familia, encargar "
    "un estudio de consultoría o hacer una selección permanente a contrarreloj. Las "
    "tres respuestas tienen un denominador común: suelen llegar tarde y sin la "
    "capacidad ejecutiva necesaria para gestionar la transición en tiempo real."
)
p = doc.add_paragraph(intro_text2)
set_body_style(p)

intro_text3 = (
    "Existe una cuarta opción que el mercado europeo utiliza de manera sistemática "
    "desde hace más de una década y que España apenas empieza a descubrir: el "
    "directivo interino especializado en transiciones directivas. No es un parche. "
    "No es una solución de emergencia. Es una herramienta de ingeniería del relevo "
    "para cuando se quiere que las cosas salgan bien desde el principio."
)
p = doc.add_paragraph(intro_text3)
set_body_style(p)

# ---------------------------------------------------------------------------
# H2 1
# ---------------------------------------------------------------------------
doc.add_heading('La gran ola de sucesión que España no puede ignorar', level=2)

p = doc.add_paragraph(
    "Los números son elocuentes. El informe del Instituto de la Empresa Familiar (2025) "
    "confirma que las empresas familiares representan el 89 % del tejido empresarial "
    "español, generan el 67 % del empleo privado —más de 6,5 millones de puestos de "
    "trabajo— y aportan el 57 % del PIB del sector privado. Son, en términos reales, "
    "la columna vertebral de la economía española."
)
set_body_style(p)

p = doc.add_paragraph(
    "El problema es que esa columna vertebral está envejeciendo. Un tercio de sus "
    "propietarios llegará a la edad de jubilación antes de 2035. Y según los datos "
    "disponibles, solo una de cada tres empresas familiares consigue superar con éxito "
    "la transición a la segunda generación. El 40 % de las familias empresarias "
    "reconoce que la sucesión es su mayor desafío, pero más del 70 % carece de un "
    "plan de sucesión formalizado."
)
set_body_style(p)

# H3 1.1
doc.add_heading('Por qué el relevo generacional es más que un cambio de nombre en la puerta del despacho', level=3)

p = doc.add_paragraph(
    "La sucesión en una empresa familiar no es un acto administrativo. Es una "
    "transformación simultánea en varias dimensiones: cambia el liderazgo, cambia la "
    "cultura de toma de decisiones, cambian las relaciones con los directivos clave, "
    "con los clientes históricos y con los proveedores que llevan décadas trabajando "
    "con el fundador. En muchos casos, cambia también la estrategia, porque la nueva "
    "generación tiene una visión diferente del negocio que quiere construir."
)
set_body_style(p)

p = doc.add_paragraph(
    "Cuando ese proceso no está bien gestionado, el resultado es predecible: fuga de "
    "talento directivo que no quiere trabajar con un sucesor sin autoridad real, "
    "pérdida de clientes que tenían una relación personal con el fundador, parálisis "
    "en la toma de decisiones mientras la familia debate quién manda, y en los peores "
    "casos, ventas forzosas en condiciones desfavorables o conflictos sucesorios que "
    "destruyen en meses lo que tardó décadas en construirse. Como señala Soly Sakal, "
    "fundador de Rhombus: «Sin una planificación clara, lo que ha tardado décadas en "
    "construirse puede desaparecer en cuestión de meses»."
)
set_body_style(p)

# H3 1.2
doc.add_heading('Lo que el mercado europeo ya sabe (y España aún no aplica)', level=3)

p = doc.add_paragraph(
    "En Francia, la sucesión del tejido empresarial familiar lleva años siendo una "
    "línea de negocio explícita para los grandes proveedores de management de "
    "transition. En Alemania, la escasez de sucesores en el Mittelstand —el tejido "
    "de empresas medianas que constituye el motor industrial alemán— ha convertido a "
    "los Interimsmanager en actores habituales y reconocidos de los procesos de relevo."
)
set_body_style(p)

p = doc.add_paragraph(
    "España replica ese patrón con un retraso estimado de 18 a 24 meses, según datos "
    "del sector recogidos por Servitalent (2025). La brecha no es de necesidad —el "
    "problema es idéntico al de los mercados de referencia—, sino de conocimiento. "
    "La mayoría de los propietarios de empresas familiares en España aún no saben "
    "que existe un profesional específico para acompañar exactamente su situación."
)
set_body_style(p)

# ---------------------------------------------------------------------------
# H2 2
# ---------------------------------------------------------------------------
doc.add_heading('Qué es exactamente un directivo interino y qué no es', level=2)

p = doc.add_paragraph(
    "Un directivo interino es un profesional senior —habitualmente con veinte o más "
    "años de experiencia ejecutiva contrastada— que asume responsabilidades de "
    "dirección de manera temporal, con autoridad real y una misión definida al inicio "
    "del mandato. No trabaja de manera indefinida. No es una contratación provisional "
    "mientras se busca al candidato definitivo. Es un experto en gestión de "
    "transiciones que entra con un objetivo claro, lo ejecuta con autonomía y sale "
    "de manera planificada, dejando a la organización en mejor posición que cuando "
    "llegó."
)
set_body_style(p)

# H3 2.1
doc.add_heading('El directivo interino no es un consultor', level=3)

p = doc.add_paragraph(
    "Esta distinción es la más importante para los propietarios de empresas familiares "
    "que han tenido experiencias previas con consultoras. El consultor diagnostica, "
    "recomienda y entrega un informe. El directivo interino ejecuta. Asume la "
    "responsabilidad operativa, toma decisiones, lidera equipos y responde por los "
    "resultados. La diferencia no es de grado: es de naturaleza."
)
set_body_style(p)

p = doc.add_paragraph(
    "En el contexto de una sucesión familiar, esto tiene un valor específico. Los "
    "propietarios no necesitan un diagnóstico más —habitualmente ya saben cuál es el "
    "problema—. Necesitan a alguien que lo gestione mientras ellos deciden, mientras "
    "el sucesor se prepara, o mientras la familia encuentra el consenso necesario "
    "para avanzar."
)
set_body_style(p)

# H3 2.2
doc.add_heading('El directivo interino no es un empleado temporal de alta dirección', level=3)

p = doc.add_paragraph(
    "Tampoco es una ETT ejecutiva. El directivo interino no se contrata a través de "
    "una empresa de trabajo temporal ni opera bajo las condiciones de un contrato "
    "laboral ordinario. Es un profesional independiente con un contrato de servicios "
    "específico, que define desde el primer día los objetivos del mandato, los plazos "
    "de ejecución y las métricas que determinarán el éxito de la misión. Esa "
    "independencia —y esa claridad de propósito— es precisamente lo que le permite "
    "operar con la velocidad y la autoridad que una transición directiva requiere."
)
set_body_style(p)

# ---------------------------------------------------------------------------
# H2 3
# ---------------------------------------------------------------------------
doc.add_heading('Los cinco momentos en que un directivo interino es la solución correcta', level=2)

p = doc.add_paragraph(
    "No todas las situaciones de sucesión son iguales. Sin embargo, hay cinco "
    "escenarios recurrentes en los que la gestión interina ofrece la respuesta más "
    "eficaz, más rápida y con mayor probabilidad de éxito:"
)
set_body_style(p)

scenarios = [
    (
        "1. El sucesor designado no está listo todavía.",
        "El hijo mayor tiene talento y voluntad, pero necesita uno o dos años de "
        "acompañamiento antes de asumir el mando con plena autoridad. Un directivo "
        "interino ocupa el puesto de dirección, estabiliza la organización y actúa "
        "simultáneamente como mentor del futuro líder. Cuando la transición se "
        "formaliza, el sucesor llega con una base sólida, no con el peso de aprender "
        "y dirigir al mismo tiempo."
    ),
    (
        "2. El fundador se retira de manera abrupta.",
        "Una enfermedad, una decisión repentina, un accidente. Cuando la salida del "
        "fundador no estaba en el plan, la empresa queda descapitalizada de dirección "
        "en cuestión de días. El directivo interino puede estar operativo en dos o "
        "tres semanas, garantizar la continuidad y dar a la familia el tiempo que "
        "necesita para tomar decisiones estratégicas sin presión."
    ),
    (
        "3. La empresa ha crecido más allá de la capacidad de gestión familiar.",
        "El negocio que empezó con diez empleados tiene ahora cien. Las herramientas "
        "de gestión que funcionaban entonces ya no son suficientes. Un directivo "
        "interino con experiencia en procesos de profesionalización puede implantar "
        "los sistemas y las estructuras que la empresa necesita antes de decidir "
        "quién la lidera de manera permanente."
    ),
    (
        "4. El sucesor ha tomado el mando pero necesita apoyo operativo sénior.",
        "El nuevo director general es joven, tiene visión estratégica, pero la "
        "operación cotidiana le desborda. Un director de operaciones o un CFO "
        "interino sénior le da la estabilidad necesaria durante los primeros doce "
        "meses, mientras el sucesor consolida su liderazgo sin que la empresa "
        "sufra las consecuencias de la curva de aprendizaje."
    ),
    (
        "5. La empresa está en proceso de venta o atracción de inversión.",
        "Cuando una empresa familiar decide incorporar un socio financiero o "
        "preparar una due diligence, la dirección neutral y profesional de un "
        "interino es un activo de credibilidad frente a los potenciales "
        "inversores. Un directivo con historial ejecutivo verificable, sin vínculos "
        "emocionales con la familia, maximiza el valor percibido de la empresa "
        "en el proceso de transacción."
    ),
]

for title, desc in scenarios:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Cm(0.5)
    run_title = p.add_run(title + ' ')
    run_title.bold = True
    run_title.font.size = Pt(11)
    run_title.font.name = 'Calibri'
    run_desc = p.add_run(desc)
    run_desc.font.size = Pt(11)
    run_desc.font.name = 'Calibri'

# CTA secundario
add_cta_box(
    doc,
    title_text="¿Te has reconocido en alguna de estas situaciones?",
    body_text=(
        "Podemos evaluar tu caso en una primera conversación confidencial y sin coste. "
        "Manager in Motion acompaña procesos de sucesión en empresas familiares con "
        "directivos interinos de perfil senior y experiencia verificada en el mercado ibérico."
    ),
    button_text="Hablar con un especialista en gestión interina"
)

# ---------------------------------------------------------------------------
# H2 4
# ---------------------------------------------------------------------------
doc.add_heading('Cómo se estructura un mandato de gestión interina en una empresa familiar', level=2)

p = doc.add_paragraph(
    "Uno de los frenos más habituales para activar una solución de gestión interina "
    "es la falta de información sobre cómo funciona el proceso. La respuesta es "
    "más sencilla de lo que parece. Un mandato bien estructurado tiene tres fases "
    "claramente diferenciadas:"
)
set_body_style(p)

phases = [
    (
        "Fase 1 — Diagnóstico y definición de misión (semanas 1-2).",
        "El proveedor de gestión interina realiza una evaluación inicial de la "
        "situación: qué se busca exactamente, cuál es el resultado esperado al final "
        "del mandato, quién reporta a quién, cuáles son las sensibilidades de la "
        "familia y qué perfil de directivo se necesita. Esta fase es crítica: una "
        "misión bien definida desde el inicio tiene una probabilidad de éxito "
        "radicalmente superior a una misión que empieza con objetivos vagos."
    ),
    (
        "Fase 2 — Selección e incorporación (semanas 2-4).",
        "Los proveedores especializados cuentan con redes de directivos disponibles "
        "que permiten activar una misión en dos a cuatro semanas desde el primer "
        "contacto. No se trata de publicar una oferta de empleo: se trata de "
        "seleccionar al profesional adecuado —con historial en situaciones similares— "
        "de una red ya disponible y cualificada."
    ),
    (
        "Fase 3 — Ejecución y transferencia (duración del mandato).",
        "El directivo interino trabaja con plena autoridad ejecutiva, reportando al "
        "Consejo de Administración o al propietario según la estructura acordada. "
        "Al final del mandato, la salida está tan planificada como la entrada: "
        "documentación de procesos, transferencia de relaciones clave, formación "
        "del sucesor y un periodo de solapamiento supervisado que asegura la "
        "continuidad de lo construido."
    ),
]

for title, desc in phases:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Cm(0.5)
    run_t = p.add_run(title + ' ')
    run_t.bold = True
    run_t.font.size = Pt(11)
    run_t.font.name = 'Calibri'
    run_d = p.add_run(desc)
    run_d.font.size = Pt(11)
    run_d.font.name = 'Calibri'

# H3 4.1
doc.add_heading('Duración típica de un mandato: ni demasiado corto ni demasiado largo', level=3)

p = doc.add_paragraph(
    "Según el Baromètre France Transition S2 2025 —la referencia estadística más "
    "completa del sector en Europa—, la duración media de los mandatos de gestión "
    "interina se sitúa en 7,5 meses, con los mandatos de transformación alcanzando "
    "los 8,3 meses de media. En contextos de sucesión familiar, los plazos tienden "
    "a extenderse hacia los 9-12 meses, porque el componente de transferencia de "
    "conocimiento, gestión de relaciones y acompañamiento del sucesor es "
    "estructuralmente más complejo que en un mandato de reestructuración financiera."
)
set_body_style(p)

p = doc.add_paragraph(
    "Este dato tiene una implicación práctica importante: una misión de gestión "
    "interina en una empresa familiar no es una solución de urgencia de tres meses. "
    "Es una inversión planificada en la estabilidad y el valor de la empresa durante "
    "el periodo más crítico de su historia. El tiempo importa, pero la calidad de "
    "la ejecución importa más."
)
set_body_style(p)

# H3 4.2
doc.add_heading('El argumento del ROI frente al coste de no actuar', level=3)

p = doc.add_paragraph(
    "La pregunta más habitual de un propietario de empresa familiar ante la "
    "posibilidad de contratar un directivo interino es: ¿cuánto cuesta? Es la "
    "pregunta equivocada. La pregunta correcta es: ¿cuánto cuesta NO hacerlo?"
)
set_body_style(p)

p = doc.add_paragraph(
    "Los datos del mercado europeo son ilustrativos. En Francia, el ROI de una "
    "misión de gestión interina bien ejecutada se estima habitualmente entre 3 y 4 "
    "veces el coste del mandato (Louis Dupont, análisis S2 2025). Las tarifas "
    "diarias de referencia en los mercados europeos avanzados se sitúan por encima "
    "de los 1.300 € en Francia y los 1.317 € en Alemania. El mercado español, más "
    "joven, opera con tarifas inferiores, lo que mejora aún más la ecuación de valor "
    "para el cliente."
)
set_body_style(p)

p = doc.add_paragraph(
    "El coste de una mala sucesión —pérdida de directivos clave, erosión de la "
    "cartera de clientes, paralización estratégica, conflictos accionariales o "
    "venta en condiciones desfavorables— es sistemáticamente mayor que el coste "
    "de contratar al profesional adecuado en el momento adecuado. Esa ecuación es "
    "la que explica por qué los mercados europeos más maduros han normalizado el "
    "uso del directivo interino en los procesos de sucesión familiar."
)
set_body_style(p)

# ---------------------------------------------------------------------------
# H2 5
# ---------------------------------------------------------------------------
doc.add_heading('Lo que las empresas familiares más avanzadas ya están haciendo', level=2)

p = doc.add_paragraph(
    "Los siguientes casos tipificados ilustran tres patrones reales de mandato de "
    "gestión interina en contextos de sucesión familiar. Los nombres son ficticios, "
    "los perfiles son representativos de situaciones reales del mercado ibérico:"
)
set_body_style(p)

cases = [
    (
        "Caso A — Empresa industrial de segunda generación (sector transformación metálica, 85 empleados).",
        "El fundador, de 67 años, tiene dos hijos en la empresa. El mayor gestiona "
        "operaciones; el menor, comercial. Ninguno de los dos tiene la visión global "
        "ni la experiencia financiera para asumir la dirección general. La empresa "
        "contrata a un CEO interino con perfil industrial y experiencia en "
        "profesionalización de pymes. Durante doce meses, el interino implanta los "
        "sistemas de reporting, reestructura el comité de dirección, incorpora un "
        "director financiero externo y acompaña al hijo mayor en el ejercicio del "
        "liderazgo. Al final del mandato, el sucesor asume la dirección con una "
        "organización profesionalizada y con credibilidad ante el equipo."
    ),
    (
        "Caso B — Empresa de distribución alimentaria (sector alimentación, 120 empleados).",
        "El fundador sufre un problema de salud grave e inesperado. No existe "
        "protocolo de sucesión. La empresa tiene contratos con grandes cadenas de "
        "distribución que vencen en seis meses. Un director general interino se "
        "incorpora en dieciséis días. Su primera tarea es estabilizar al equipo "
        "directivo —tres de sus cuatro directores estaban considerando marcharse—, "
        "renovar los contratos con las cadenas y preparar a la hija del fundador, "
        "que tiene veintiocho años y un MBA, para asumir la dirección en el "
        "horizonte de nueve meses. La empresa no pierde ni un contrato. La "
        "transición al sucesor se produce en los plazos previstos."
    ),
    (
        "Caso C — Empresa de servicios profesionales (sector ingeniería, 45 empleados).",
        "La empresa familiar recibe una oferta de adquisición parcial de un fondo "
        "de private equity. Los propietarios quieren maximizar el valor antes de "
        "la due diligence, pero la dirección financiera interna no tiene el nivel "
        "requerido para las negociaciones. Un CFO interino con experiencia en "
        "transacciones M&A se incorpora durante ocho meses. Implanta el sistema "
        "de control de gestión que el fondo exige, lidera la due diligence y "
        "asesora a la familia en la negociación de la valoración. La operación "
        "se cierra con una valoración un 22 % superior a la oferta inicial."
    ),
]

for title, desc in cases:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Cm(0.5)
    run_t = p.add_run(title + ' ')
    run_t.bold = True
    run_t.font.size = Pt(11)
    run_t.font.name = 'Calibri'
    run_d = p.add_run(desc)
    run_d.font.size = Pt(11)
    run_d.font.name = 'Calibri'

# ---------------------------------------------------------------------------
# H2 6
# ---------------------------------------------------------------------------
doc.add_heading('Cómo elegir al proveedor de gestión interina adecuado para tu empresa familiar', level=2)

p = doc.add_paragraph(
    "No todos los proveedores de gestión interina son iguales. En un mercado en "
    "proceso de maduración como el español, la calidad varía significativamente. "
    "Estos son los criterios que distinguen a un proveedor adecuado para un proceso "
    "de sucesión familiar:"
)
set_body_style(p)

criteria = [
    "Experiencia documentada en contextos familiares. No basta con haber gestionado "
    "vacantes directivas en grandes corporaciones. La dinámica familia-empresa tiene "
    "sus propias reglas, y el directivo interino que no las conoce puede generar más "
    "conflicto que valor.",

    "Metodología de diagnóstico inicial. Un buen proveedor no propone un candidato "
    "antes de entender la situación. La primera conversación debe ser sobre el "
    "problema, no sobre el catálogo de perfiles disponibles.",

    "Garantías de confidencialidad. En una empresa familiar, la información sobre "
    "la sucesión es especialmente sensible. El proveedor debe tener protocolos "
    "explícitos de confidencialidad para el proceso de selección y para el "
    "desarrollo del mandato.",

    "Portfolio de mandatos similares. Preguntar por casos anteriores en empresas "
    "del mismo sector, tamaño y tipo de situación. Un proveedor con experiencia "
    "real en sucesiones familiares puede dar referencias verificables.",

    "Capacidad de gestión de la salida. El final del mandato es tan importante como "
    "el inicio. Un proveedor solvente tiene metodología específica para planificar "
    "la transferencia de conocimiento y la salida del directivo sin generar un "
    "nuevo vacío en la organización.",
]

for idx, criterion in enumerate(criteria, 1):
    p = doc.add_paragraph(style='List Number')
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(criterion)
    run.font.size = Pt(11)
    run.font.name = 'Calibri'

# ---------------------------------------------------------------------------
# H2 7 — CONCLUSIÓN
# ---------------------------------------------------------------------------
doc.add_heading('El relevo generacional como momento de transformación, no de crisis', level=2)

p = doc.add_paragraph(
    "La sucesión en una empresa familiar no tiene por qué ser un momento de riesgo. "
    "Puede ser el mayor salto de valor que la empresa ha experimentado en décadas: "
    "el momento en que se profesionaliza la dirección, se modernizan los procesos, "
    "se construye la estructura que va a sostener el crecimiento de la siguiente "
    "generación."
)
set_body_style(p)

p = doc.add_paragraph(
    "La diferencia entre una sucesión que destruye valor y una sucesión que lo "
    "multiplica reside, en muchos casos, en una sola decisión: activar el apoyo "
    "ejecutivo adecuado en el momento adecuado. El directivo interino especializado "
    "en transiciones familiares no resuelve el problema por la familia —esa es "
    "siempre una decisión que corresponde a la familia—, pero crea las condiciones "
    "para que esa decisión pueda tomarse desde la fortaleza y no desde la urgencia."
)
set_body_style(p)

p = doc.add_paragraph(
    "El mercado europeo ya lo sabe. España está aprendiendo. Las empresas familiares "
    "que actúen primero —que planifiquen el relevo con la misma seriedad con que "
    "planifican su expansión comercial o su estrategia de producto— serán las que "
    "lideren la siguiente generación de su sector."
)
set_body_style(p)

add_horizontal_rule(doc)

# ---------------------------------------------------------------------------
# FAQ
# ---------------------------------------------------------------------------
doc.add_heading('Preguntas frecuentes sobre el directivo interino en la empresa familiar', level=2)

faqs = [
    (
        "¿Puede un directivo interino asumir el cargo de CEO en una empresa familiar durante la transición?",
        "Sí. Un directivo interino puede asumir cualquier posición de la alta dirección "
        "—CEO, COO, CFO, Director General— de manera temporal y con plena autoridad "
        "ejecutiva. El mandato se define al inicio del proceso, con objetivos claros, "
        "hitos verificables y una fecha de salida planificada desde el primer día."
    ),
    (
        "¿Cuánto tiempo tarda en incorporarse un directivo interino desde que se toma la decisión?",
        "En condiciones normales, el proceso de selección e incorporación dura entre "
        "dos y cuatro semanas. Los proveedores especializados tienen redes de "
        "directivos disponibles que permiten activar una misión con rapidez cuando la "
        "urgencia lo requiere. En situaciones de emergencia, los plazos pueden "
        "acortarse a menos de dos semanas."
    ),
    (
        "¿El directivo interino puede generar conflicto con los miembros de la familia que siguen en la empresa?",
        "Es uno de los riesgos más citados, pero también uno de los más gestionables. "
        "Un directivo interino con experiencia en contextos familiares sabe trabajar "
        "con la dualidad familia-empresa, establecer límites claros y ganarse la "
        "confianza de las distintas partes. La clave está en la definición precisa "
        "del mandato desde el inicio y en el respaldo explícito del propietario o "
        "del Consejo de Administración."
    ),
    (
        "¿Es la gestión interina adecuada para empresas pequeñas con menos de 50 empleados?",
        "Depende de la complejidad del relevo y del contexto específico. En empresas "
        "de entre 20 y 50 empleados, puede ser más apropiado un directivo interino a "
        "tiempo parcial —modalidad creciente en el mercado español, donde el 75 % de "
        "las misiones ya se ejecutan en formato part-time (Servitalent, 2025)— que un "
        "mandato a jornada completa. La solución debe dimensionarse al problema real."
    ),
    (
        "¿Qué ocurre cuando el directivo interino termina su mandato? ¿Se pierde lo que ha construido?",
        "Una buena gestión del final del mandato incluye un plan de transferencia "
        "explícito: documentación de procesos, traspaso de relaciones clave, "
        "formación del sucesor y un periodo de solapamiento supervisado. Los "
        "mandatos bien diseñados producen resultados que perduran más allá de la "
        "presencia del interino. La transferencia de valor es parte del servicio, "
        "no un extra opcional."
    ),
    (
        "¿Puede el directivo interino convertirse en el director permanente si la empresa queda satisfecha?",
        "Técnicamente sí, pero no es el objetivo ni la expectativa del modelo. El "
        "management interino funciona precisamente porque el profesional es temporal "
        "por naturaleza: su independencia y su foco en el resultado se apoyan en "
        "saber que la misión tiene fin. Convertir el mandato en permanente puede ser "
        "contraproducente para ambas partes, aunque hay casos (cerca del 10 % según "
        "datos de GH Partners, 2025) en que el proceso concluye con una incorporación "
        "permanente por decisión consensuada."
    ),
]

for question, answer in faqs:
    p_q = doc.add_paragraph()
    p_q.paragraph_format.space_before = Pt(10)
    p_q.paragraph_format.space_after = Pt(2)
    r = p_q.add_run(question)
    r.bold = True
    r.font.size = Pt(11)
    r.font.name = 'Calibri'

    p_a = doc.add_paragraph(answer)
    p_a.paragraph_format.space_before = Pt(2)
    p_a.paragraph_format.space_after = Pt(6)
    for run in p_a.runs:
        run.font.size = Pt(11)
        run.font.name = 'Calibri'

add_horizontal_rule(doc)

# ---------------------------------------------------------------------------
# CTA PRINCIPAL
# ---------------------------------------------------------------------------
add_cta_box(
    doc,
    title_text="¿Estás planificando el relevo de tu empresa familiar o afrontando una transición directiva urgente?",
    body_text=(
        "Manager in Motion acompaña procesos de sucesión en empresas familiares con "
        "directivos interinos de perfil senior y experiencia verificada en el mercado "
        "ibérico. Primera conversación confidencial y sin compromiso."
    ),
    button_text="Solicitar diagnóstico inicial gratuito"
)

# ---------------------------------------------------------------------------
# FUENTES
# ---------------------------------------------------------------------------
doc.add_paragraph()
p_sources = doc.add_paragraph()
r = p_sources.add_run('Fuentes y referencias')
r.bold = True
r.font.size = Pt(10)
r.font.name = 'Calibri'
r.font.color.rgb = RGBColor(0x77, 0x77, 0x77)

sources_list = [
    "Instituto de la Empresa Familiar — Informe «Relevancia y supervivencia de la empresa familiar. La Empresa Familiar en España 2025» — iefamiliar.com",
    "Baromètre France Transition S2 2025 — análisis Louis Dupont (publicado abril 2026) — louis-dupont.com",
    "GH Partners — «Le marché français du management de transition — chiffres clés, tendances 2025-2026» — ghpartners.fr",
    "AIME — Asociación Interim Management España — VII Congreso Nacional de Interim Management (12 de mayo de 2026) — interimspain.org",
    "Servitalent — «Sucesión en empresa familiar. Interim Management para asegurar valor» (2025) — servitalent.com",
    "EIM España — «Empresa familiar y relevo generacional: el valor que aporta el Interim Management» — eim.com",
    "Rhombus / RRHH Digital — «El relevo generacional, principal riesgo para la continuidad de miles de pymes familiares en España» (2025-2026)",
]

for source in sources_list:
    p = doc.add_paragraph(source, style='List Bullet')
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    for run in p.runs:
        run.font.size = Pt(9)
        run.font.name = 'Calibri'
        run.font.color.rgb = RGBColor(0x77, 0x77, 0x77)

# ---------------------------------------------------------------------------
# Save
# ---------------------------------------------------------------------------
os.makedirs(os.path.dirname(OUTPUT_PATH), exist_ok=True)
doc.save(OUTPUT_PATH)
print(f"Documento guardado en: {OUTPUT_PATH}")
