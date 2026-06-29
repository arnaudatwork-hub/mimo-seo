#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Generador de artículo SEO — Director Financiero Interino
Manager in Motion — 2026-06-29
"""

from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

# ---------------------------------------------------------------------------
# Contenido del artículo
# ---------------------------------------------------------------------------

TITULO_SEO = "Director Financiero Interino: cuándo contratarlo y por qué en 2026 | Manager in Motion"
META_DESC = ("La demanda de CFO interino ha crecido un 103% en Europa. Descubre cuándo tu empresa "
             "necesita un director financiero de transición, qué perfil buscar y cómo estructurar "
             "la misión. Guía práctica para CEOs y consejos de administración.")
SLUG = "/director-financiero-interino"
H1 = "Director financiero interino: cuándo lo necesita tu empresa y cómo aprovecharlo"

INTRO = (
    "El pasado ejercicio, la demanda de CFOs interinos creció un 103% en Europa. No es un dato "
    "coyuntural: es la señal de que la función financiera ha dejado de ser la última en adoptar el "
    "modelo de liderazgo temporal y se ha convertido en la primera. En España, el 98% de los "
    "expertos y directivos identifica la escasez de talento directivo como el principal freno a la "
    "productividad empresarial. Cuando esa escasez afecta a la dirección financiera, las "
    "consecuencias son inmediatas y costosas. El director financiero interino es la respuesta que "
    "muchas empresas españolas aún no han considerado, pero que sus homólogas europeas llevan años "
    "aplicando con resultados medibles."
)

H2_1 = "La escasez de CFOs cualificados: un problema estructural con solución probada"

CUERPO_1 = (
    "En el Reino Unido, el 51% de todas las solicitudes de liderazgo interino corresponde a la "
    "función financiera. Es el perfil más demandado también en Francia y Alemania, donde el mercado "
    "de interim management alcanzó en 2026 un volumen récord de 2.700 millones de euros con una "
    "tarifa media de 1.317 euros por jornada. España sigue esta misma curva con un desfase estimado "
    "de 12 a 18 meses.\n\n"
    "La escasez no es solo cuantitativa. Los perfiles financieros con experiencia en "
    "reestructuración, M&A o transformación digital son escasos, tienen ciclos de selección de "
    "cuatro a nueve meses y generan costes de incorporación que raramente se reflejan en el "
    "presupuesto inicial. Cuando la vacante es urgente, o cuando el contexto exige un perfil que la "
    "empresa no puede o no quiere contratar de forma permanente, el modelo de interim management "
    "resuelve el problema en un plazo que un proceso convencional no puede igualar.\n\n"
    "El 44% de las empresas españolas ha redefinido sus planes de crecimiento en 2026. Muchas de "
    "esas revisiones implican decisiones financieras complejas para las que la estructura actual no "
    "está preparada. La respuesta no es esperar a encontrar al candidato permanente ideal: es cubrir "
    "la función con un interim manager que ejecute desde el primer día."
)

H2_2 = "Qué es un director financiero interino y en qué situaciones encaja"

CUERPO_2_P1 = (
    "Un director financiero interino no es un consultor que entrega un informe y se marcha. Es un "
    "directivo con autoridad de línea: gestiona el equipo financiero, toma decisiones vinculantes, "
    "interactúa con bancos, inversores y el consejo de administración, y responde ante el CEO con "
    "la misma responsabilidad que un CFO en plantilla. La diferencia está en el contrato, no en el "
    "mandato."
)

CUERPO_2_P2 = "El interim manager financiero encaja con precisión en seis situaciones:"

SITUACIONES = [
    ("Vacante urgente por salida inesperada",
     "del director financiero, sin margen para un proceso de selección de meses."),
    ("Crecimiento acelerado",
     "que desborda la estructura financiera actual: empresas que escalan y necesitan un CFO que "
     "construya la función desde cero."),
    ("Operaciones de M&A",
     "ya sea en la fase de due diligence, integración post-adquisición o preparación para el exit. "
     "El directivo interino en operaciones de private equity aporta experiencia que difícilmente se "
     "encuentra en el mercado permanente."),
    ("Reestructuración financiera o turnaround",
     "donde la neutralidad y la experiencia del interim manager son activos críticos."),
    ("Transformación digital de la función financiera",
     "cuando la empresa necesita implantar un ERP, automatizar procesos o rediseñar el modelo de "
     "reporting."),
    ("Profesionalización de la empresa familiar",
     "especialmente en la fase previa a la sucesión del fundador o en la preparación de la empresa "
     "para acceder a financiación bancaria o inversión externa."),
]

CUERPO_2_P3 = (
    "El directivo financiero de transición actúa durante el tiempo estrictamente necesario, con un "
    "mandato claro desde el inicio y un plan de salida definido. Las misiones duran habitualmente "
    "entre tres y doce meses, con mayor concentración en el tramo de seis a nueve meses."
)

CUERPO_2_P4 = (
    "En materia de coste, las tarifas de referencia en España oscilan entre 800 y 1.500 euros por "
    "jornada. Comparado con el coste total de un CFO permanente, que incluye salario bruto, cargas "
    "sociales, beneficios, coste de selección y eventual indemnización por despido, el CFO interino "
    "resulta financieramente competitivo en misiones de hasta doce meses, y estructuralmente "
    "superior cuando la necesidad tiene un horizonte definido."
)

CUERPO_2_CTA_INLINE = (
    "Si se reconoce alguna de estas situaciones en la empresa, una primera conversación con un "
    "especialista en interim management no compromete a nada. [enlace: hablar con un especialista → /contacto]"
)

H2_3 = "Cómo funciona una misión: del primer contacto al traspaso"

CUERPO_3 = (
    "Una misión de director financiero interino bien estructurada sigue cuatro fases. La primera es "
    "el diagnóstico: antes de incorporar al interim manager, se define el mandato con precisión. Qué "
    "debe resolver, en qué plazo, con qué recursos y con qué indicadores de éxito. Esta fase dura "
    "entre dos y cinco días y es determinante para el resultado.\n\n"
    "La segunda fase es el arranque. Las primeras dos a cuatro semanas son de integración activa: el "
    "interim manager analiza la situación financiera real, identifica los riesgos inmediatos, se "
    "presenta ante los interlocutores clave y establece una hoja de ruta. No hay período de "
    "cortesía: el trabajo comienza desde el primer día.\n\n"
    "La tercera fase es la ejecución: el CFO interino opera con plena autoridad de línea, lidera el "
    "equipo, ejecuta el plan y reporta al CEO o al consejo con la misma cadencia que un director "
    "financiero permanente. La diferencia con un consultor es visible aquí: no hay recomendaciones "
    "pendientes de implementación, hay resultados.\n\n"
    "La cuarta fase, y con frecuencia la más descuidada, es la transferencia. Una salida bien "
    "gestionada incluye la documentación de los procesos, la formación del sucesor, ya sea interno "
    "o externo, y una transición que no genere dependencia. El interim manager que deja la empresa "
    "mejor de como la encontró, sin crear una nueva vacante crítica, es el que ha cumplido su "
    "mandato.\n\n"
    "Manager in Motion despliega perfiles financieros ejecutivos en menos de 72 horas desde la "
    "decisión. La selección de un director financiero interino [enlace: /servicios/cfo-interino] "
    "no requiere cuatro meses: requiere un briefing claro y un proveedor con el talento disponible. "
    "Para empresas que gestionan situaciones urgentes, esa diferencia de tiempo es, con frecuencia, "
    "la diferencia entre controlar la situación y perder el margen de maniobra."
)

# ---------------------------------------------------------------------------
# FAQ (3 preguntas máximo según instrucciones del sistema)
# ---------------------------------------------------------------------------

FAQS = [
    (
        "¿Qué es un director financiero interino?",
        "Un director financiero interino es un directivo con plena autoridad ejecutiva que se "
        "incorpora a la empresa por un período definido para liderar la función financiera. No es "
        "un consultor ni un asesor externo: gestiona equipos, toma decisiones vinculantes, firma "
        "con entidades financieras y responde ante el CEO o el consejo de administración con la "
        "misma responsabilidad que un CFO en plantilla. La diferencia fundamental respecto a otros "
        "modelos es que ejecuta, no recomienda."
    ),
    (
        "¿Cuánto cuesta un director financiero interino en España?",
        "Las tarifas de referencia en España oscilan entre 800 y 1.500 euros por jornada, en "
        "función del perfil, la complejidad del mandato y el sector. Frente a un CFO permanente, "
        "cuyo coste total (salario bruto, cargas sociales, beneficios, selección e indemnización "
        "potencial) puede superar los 200.000 euros anuales, el CFO interino resulta "
        "financieramente competitivo en misiones de hasta doce meses y especialmente eficiente "
        "cuando la necesidad tiene un horizonte temporal definido."
    ),
    (
        "¿En qué se diferencia un CFO interino de un consultor financiero?",
        "La diferencia es de autoridad y de responsabilidad. El consultor financiero analiza, "
        "diagnostica y entrega recomendaciones; la implementación recae en la empresa. El interim "
        "manager financiero implementa directamente: gestiona el equipo, opera los sistemas, "
        "representa a la empresa ante bancos e inversores y toma decisiones de línea. Cuando el "
        "problema requiere acción, no solo diagnóstico, el modelo de interim management es el "
        "adecuado."
    ),
]

# ---------------------------------------------------------------------------
# CTA
# ---------------------------------------------------------------------------

CTA_TITULO = "¿Tu empresa necesita un director financiero interino?"
CTA_CUERPO = (
    "En Manager in Motion desplegamos el perfil adecuado en menos de 72 horas. Cuéntanos tu "
    "situación y te presentamos una propuesta sin compromiso."
)
CTA_BOTON = "Solicitar un CFO interino [enlace: /contacto]"

# ---------------------------------------------------------------------------
# Enlazado interno
# ---------------------------------------------------------------------------

ENLAZADO_INTERNO = [
    '**interim management** [enlace: /servicios/interim-management] — primera mención explicativa del concepto (sección H2-1)',
    '**director financiero interino** [enlace: /servicios/cfo-interino] — sección H2-3 (metodología) y CTA',
    '**directivo interino en operaciones de private equity** [enlace: /blog/interim-management-private-equity-espana] — situación 3 (M&A)',
    '**director general interino** [enlace: /servicios/ceo-interino] — sección de posicionamiento / CTA',
    '**solicitar un CFO interino** [enlace: /contacto] — CTA principal',
]

# ---------------------------------------------------------------------------
# Fuentes consultadas
# ---------------------------------------------------------------------------

FUENTES = [
    "Heidrick & Struggles — Talent Lens Survey 2026. Dato: crecimiento del 103% interanual en la "
    "demanda de CFOs interinos en Europa; el 51% de las solicitudes de liderazgo interino en el "
    "Reino Unido corresponde a la función financiera. URL de referencia: heidrick.com",
    "PwC España — Informe 2026 sobre escasez de talento directivo. Dato: el 98% de los expertos y "
    "directivos en España identifica la escasez de talento directivo como el primer freno a la "
    "productividad empresarial. URL de referencia: pwc.es",
    "EY España — Perspectivas del mercado de trabajo 2026. Dato: vulnerabilidad de las pymes ante "
    "el incremento de costes laborales unitarios y escasez de talento cualificado. URL de referencia: "
    "ey.com/es_ES",
    "KPMG España / CEOE — Perspectivas 2026. Dato: el 44% de las empresas españolas ha redefinido "
    "sus planes de crecimiento en 2026. URL de referencia: kpmg.com/es",
    "DDIM — Marktstudie 2026 (Dachgesellschaft Deutsches Interim Management). Dato: el mercado "
    "alemán de interim management alcanzó un volumen récord de 2.700 millones de euros en 2026 con "
    "una tarifa media de 1.317 euros/día. URL de referencia: ddim.de",
]

NOTAS_REVISION = (
    "1. Verificar que las URLs de enlazado interno corresponden a páginas publicadas en "
    "managerinmotion.eu antes de publicar.\n"
    "2. Confirmar con el equipo de datos que las cifras del Heidrick & Struggles Talent Lens "
    "Survey 2026, PwC España 2026, KPMG/CEOE 2026 y DDIM Marktstudie 2026 son las versiones "
    "finales publicadas (no borradores o estimaciones preliminares).\n"
    "3. Añadir schema markup FAQPage en el CMS para las tres preguntas de la sección FAQ.\n"
    "4. Implementar Article schema en el head de la página con los metadatos indicados en el brief.\n"
    "5. El CTA secundario inline (tras la sección de situaciones) debe vincularse a /contacto con "
    "seguimiento UTM: ?utm_source=blog&utm_medium=organic&utm_campaign=cfo-interino-2026."
)


# ---------------------------------------------------------------------------
# Helpers de formato
# ---------------------------------------------------------------------------

def set_heading_style(paragraph, level, text, color_hex):
    """Aplica estilo de título con color personalizado."""
    paragraph.style = f"Heading {level}"
    run = paragraph.runs[0] if paragraph.runs else paragraph.add_run(text)
    if not paragraph.runs:
        pass
    run.font.color.rgb = RGBColor(
        int(color_hex[0:2], 16),
        int(color_hex[2:4], 16),
        int(color_hex[4:6], 16),
    )


def add_left_border(paragraph):
    """Añade borde izquierdo a un párrafo (estilo cita)."""
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "24")
    left.set(qn("w:space"), "4")
    left.set(qn("w:color"), "2E75B6")
    pBdr.append(left)
    pPr.append(pBdr)


def add_paragraph(doc, text, bold_part=None, italic=False, space_after=160):
    """Añade un párrafo Normal con formato opcional."""
    p = doc.add_paragraph()
    p.style = "Normal"
    p.paragraph_format.space_after = Pt(space_after / 20)
    if bold_part:
        run_bold = p.add_run(bold_part)
        run_bold.bold = True
        remainder = text[len(bold_part):]
        run_rest = p.add_run(remainder)
        if italic:
            run_rest.italic = True
    else:
        run = p.add_run(text)
        if italic:
            run.italic = True
    return p


def add_multiline(doc, text, italic=False):
    """Divide texto en párrafos por doble salto de línea."""
    blocks = text.split("\n\n")
    for block in blocks:
        block = block.strip()
        if block:
            p = doc.add_paragraph()
            p.style = "Normal"
            run = p.add_run(block)
            if italic:
                run.italic = True
            p.paragraph_format.space_after = Pt(8)


# ---------------------------------------------------------------------------
# Construcción del documento
# ---------------------------------------------------------------------------

doc = Document()

# --- Márgenes A4 ---
section = doc.sections[0]
section.page_width = Cm(21)
section.page_height = Cm(29.7)
section.left_margin = Cm(2.54)
section.right_margin = Cm(2.54)
section.top_margin = Cm(2.54)
section.bottom_margin = Cm(2.54)

# --- Fuente base ---
style_normal = doc.styles["Normal"]
style_normal.font.name = "Arial"
style_normal.font.size = Pt(11)

# --- Estilos de título ---
for level, size, color in [
    (1, 22, "1F3864"),
    (2, 16, "2E75B6"),
    (3, 13, "404040"),
]:
    st = doc.styles[f"Heading {level}"]
    st.font.name = "Arial"
    st.font.size = Pt(size)
    st.font.bold = True
    st.font.color.rgb = RGBColor(
        int(color[0:2], 16), int(color[2:4], 16), int(color[4:6], 16)
    )
    st.paragraph_format.space_before = Pt(level == 1 and 0 or 14)
    st.paragraph_format.space_after = Pt(6)

# ===========================
# BLOQUE DE METADATOS SEO
# ===========================
meta_items = [
    ("Título SEO: ", TITULO_SEO),
    ("Meta descripción: ", META_DESC),
    ("URL slug: ", SLUG),
]
for label, value in meta_items:
    p = doc.add_paragraph()
    p.style = "Normal"
    p.paragraph_format.space_after = Pt(4)
    r_label = p.add_run(label)
    r_label.bold = True
    r_label.font.color.rgb = RGBColor(0x40, 0x40, 0x40)
    r_value = p.add_run(value)
    r_value.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

doc.add_paragraph()  # separador visual

# ===========================
# H1
# ===========================
h1 = doc.add_heading(H1, level=1)

# ===========================
# INTRODUCCIÓN
# ===========================
add_paragraph(doc, INTRO, space_after=200)

# ===========================
# H2 — 1
# ===========================
doc.add_heading(H2_1, level=2)
add_multiline(doc, CUERPO_1)
doc.add_paragraph()

# ===========================
# H2 — 2
# ===========================
doc.add_heading(H2_2, level=2)
add_paragraph(doc, CUERPO_2_P1)
add_paragraph(doc, CUERPO_2_P2)

# Lista de situaciones con número en negrita
for i, (titulo, desc) in enumerate(SITUACIONES, 1):
    p = doc.add_paragraph()
    p.style = "Normal"
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_after = Pt(5)
    run_num = p.add_run(f"{i}. ")
    run_num.bold = True
    run_titulo = p.add_run(titulo + ". ")
    run_titulo.bold = True
    run_desc = p.add_run(desc)

doc.add_paragraph()
add_paragraph(doc, CUERPO_2_P3)
add_paragraph(doc, CUERPO_2_P4)

# CTA inline
p_cta_inline = doc.add_paragraph()
p_cta_inline.style = "Normal"
p_cta_inline.paragraph_format.space_before = Pt(10)
p_cta_inline.paragraph_format.space_after = Pt(10)
add_left_border(p_cta_inline)
run_cta = p_cta_inline.add_run(CUERPO_2_CTA_INLINE)
run_cta.italic = True
run_cta.font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)

doc.add_paragraph()

# ===========================
# H2 — 3
# ===========================
doc.add_heading(H2_3, level=2)
add_multiline(doc, CUERPO_3)
doc.add_paragraph()

# ===========================
# SEPARADOR EDITORIAL
# ===========================
sep = doc.add_paragraph("─" * 60)
sep.style = "Normal"
sep.paragraph_format.space_before = Pt(20)
sep.paragraph_format.space_after = Pt(4)
sep.runs[0].font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)

# ===========================
# SECCIÓN FAQ
# ===========================
h_faq = doc.add_heading("Preguntas Frecuentes", level=2)

for pregunta, respuesta in FAQS:
    h_q = doc.add_heading(pregunta, level=3)
    h_q.runs[0].bold = True
    add_paragraph(doc, respuesta, space_after=120)

doc.add_paragraph()

# ===========================
# CTA PRINCIPAL
# ===========================
h_cta = doc.add_heading("¿Tu empresa necesita un director financiero interino?", level=2)

p_cta_body = doc.add_paragraph()
p_cta_body.style = "Normal"
p_cta_body.paragraph_format.space_after = Pt(6)
add_left_border(p_cta_body)
run_cta_body = p_cta_body.add_run(CTA_CUERPO)
run_cta_body.italic = True

p_cta_btn = doc.add_paragraph()
p_cta_btn.style = "Normal"
p_cta_btn.paragraph_format.space_after = Pt(6)
add_left_border(p_cta_btn)
run_btn = p_cta_btn.add_run(CTA_BOTON)
run_btn.bold = True
run_btn.italic = True
run_btn.font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)

doc.add_paragraph()

# ===========================
# SUGERENCIAS DE ENLAZADO INTERNO
# ===========================
h_enlaces = doc.add_heading("Sugerencias de Enlazado Interno", level=2)

for item in ENLAZADO_INTERNO:
    p = doc.add_paragraph()
    p.style = "Normal"
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(f"• {item}")
    run.italic = True

doc.add_paragraph()

# ===========================
# SEPARADOR EDITORIAL FINAL
# ===========================
sep2 = doc.add_paragraph("─" * 60)
sep2.style = "Normal"
sep2.paragraph_format.space_before = Pt(20)
sep2.paragraph_format.space_after = Pt(4)
sep2.runs[0].font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)

p_nota_ed = doc.add_paragraph()
p_nota_ed.style = "Normal"
p_nota_ed.paragraph_format.space_after = Pt(8)
run_nota = p_nota_ed.add_run("FUERA DEL ARTÍCULO — Solo para uso editorial")
run_nota.bold = True
run_nota.italic = True
run_nota.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

# ===========================
# FUENTES CONSULTADAS
# ===========================
doc.add_heading("Fuentes Consultadas", level=2)

for i, fuente in enumerate(FUENTES, 1):
    p = doc.add_paragraph()
    p.style = "Normal"
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_after = Pt(4)
    run_num = p.add_run(f"[{i}] ")
    run_num.bold = True
    run_num.italic = True
    run_num.font.color.rgb = RGBColor(0x60, 0x60, 0x60)
    run_text = p.add_run(fuente)
    run_text.italic = True
    run_text.font.color.rgb = RGBColor(0x60, 0x60, 0x60)

doc.add_paragraph()

# ===========================
# NOTAS DE REVISIÓN FINAL
# ===========================
doc.add_heading("Notas de Revisión Final", level=2)

for linea in NOTAS_REVISION.split("\n"):
    linea = linea.strip()
    if linea:
        p = doc.add_paragraph()
        p.style = "Normal"
        p.paragraph_format.left_indent = Cm(0.5)
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(linea)
        run.italic = True
        run.font.color.rgb = RGBColor(0x60, 0x60, 0x60)

# ===========================
# GUARDAR
# ===========================
output_path = "/home/user/mimo-seo/output/agent3_article.docx"
doc.save(output_path)
print(f"Archivo generado correctamente: {output_path}")
