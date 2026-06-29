#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Generador de artículo SEO — Manager in Motion
CFO interino en España
Fecha: 29 de junio de 2026
"""

from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

doc = Document()

# ── Configuración de página A4 ──────────────────────────────────────────────
section = doc.sections[0]
section.page_width  = Cm(21)
section.page_height = Cm(29.7)
section.left_margin   = Cm(2.54)
section.right_margin  = Cm(2.54)
section.top_margin    = Cm(2.54)
section.bottom_margin = Cm(2.54)

# ── Helpers ──────────────────────────────────────────────────────────────────
def set_heading_style(paragraph, level, text, color_hex):
    """Aplica estilo Heading N y color personalizado."""
    paragraph.style = f'Heading {level}'
    run = paragraph.add_run(text)
    run.font.color.rgb = RGBColor(
        int(color_hex[0:2], 16),
        int(color_hex[2:4], 16),
        int(color_hex[4:6], 16)
    )
    return run

def add_h1(doc, text):
    p = doc.add_paragraph()
    p.style = 'Heading 1'
    run = p.add_run(text)
    run.font.size = Pt(22)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(12)
    return p

def add_h2(doc, text):
    p = doc.add_paragraph()
    p.style = 'Heading 2'
    run = p.add_run(text)
    run.font.size = Pt(16)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(8)
    return p

def add_h3(doc, text):
    p = doc.add_paragraph()
    p.style = 'Heading 3'
    run = p.add_run(text)
    run.font.size = Pt(13)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x40, 0x40, 0x40)
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(6)
    return p

def add_body(doc, text):
    p = doc.add_paragraph()
    p.style = 'Normal'
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.name = 'Arial'
    p.paragraph_format.space_after = Pt(8)
    return p

def add_body_bold(doc, parts):
    """parts = lista de tuplas (texto, bold)"""
    p = doc.add_paragraph()
    p.style = 'Normal'
    for text, bold in parts:
        run = p.add_run(text)
        run.bold = bold
        run.font.size = Pt(11)
        run.font.name = 'Arial'
    p.paragraph_format.space_after = Pt(8)
    return p

def add_separator(doc):
    p = doc.add_paragraph()
    run = p.add_run('─' * 72)
    run.font.color.rgb = RGBColor(0xBB, 0xBB, 0xBB)
    run.font.size = Pt(9)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(12)
    return p

def add_meta_table(doc, titulo_seo, meta_desc, slug):
    """Tabla de metadatos SEO al inicio del documento."""
    table = doc.add_table(rows=3, cols=2)
    table.style = 'Table Grid'
    cells = [
        ('Título SEO', titulo_seo),
        ('Meta descripción', meta_desc),
        ('URL slug', slug),
    ]
    for i, (label, value) in enumerate(cells):
        row = table.rows[i]
        cell_label = row.cells[0]
        cell_value = row.cells[1]
        # Label
        p_label = cell_label.paragraphs[0]
        run_label = p_label.add_run(label)
        run_label.bold = True
        run_label.font.size = Pt(10)
        run_label.font.name = 'Arial'
        # Value
        p_value = cell_value.paragraphs[0]
        run_value = p_value.add_run(value)
        run_value.font.size = Pt(10)
        run_value.font.name = 'Arial'
    doc.add_paragraph()  # espacio tras la tabla
    return table

# ════════════════════════════════════════════════════════════════════════════
# BLOQUE 1 — METADATOS SEO
# ════════════════════════════════════════════════════════════════════════════
p_meta_title = doc.add_paragraph()
run_mt = p_meta_title.add_run('METADATOS SEO')
run_mt.bold = True
run_mt.font.size = Pt(10)
run_mt.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
run_mt.font.name = 'Arial'
p_meta_title.paragraph_format.space_after = Pt(4)

add_meta_table(
    doc,
    titulo_seo='CFO interino en España: cuándo y cómo contratarlo',
    meta_desc='Vacante urgente, reestructuración o crecimiento acelerado: descubre cuándo el CFO interino es la decisión más rentable para tu empresa en España.',
    slug='/blog/cfo-interino-espana'
)

doc.add_paragraph()  # separador visual

# ════════════════════════════════════════════════════════════════════════════
# BLOQUE 2 — H1
# ════════════════════════════════════════════════════════════════════════════
add_h1(doc, 'CFO interino en España: cuándo la dirección financiera temporal es la decisión más inteligente')

# ════════════════════════════════════════════════════════════════════════════
# BLOQUE 3 — INTRODUCCIÓN
# ════════════════════════════════════════════════════════════════════════════
add_body(doc,
    'La demanda de CFOs interinos ha crecido un 103% interanual en Europa. No es una tendencia marginal ni un dato de nicho: '
    'es la señal más clara de que la función financiera ejecutiva se está replanteando en profundidad. En paralelo, el 98% de '
    'los directivos españoles identifica la escasez de talento directivo como el principal freno a la productividad de sus '
    'empresas. La pregunta que surge de este cruce no es si el mercado necesita CFOs interinos: es cuándo el CFO interino deja '
    'de ser una solución de emergencia para convertirse en la decisión estratégica más inteligente.'
)

# ════════════════════════════════════════════════════════════════════════════
# BLOQUE 4 — H2 1
# ════════════════════════════════════════════════════════════════════════════
add_h2(doc, 'Qué es un CFO interino y por qué no es lo mismo que un consultor financiero')

add_body(doc,
    'Un interim manager financiero se incorpora a la empresa con autoridad ejecutiva real. Ocupa el cargo de director '
    'financiero, forma parte del comité de dirección, responde ante el CEO o el consejo de administración, y toma decisiones '
    'con efecto directo sobre los resultados. No entrega un informe: ejecuta. No recomienda: es responsable.'
)

add_body(doc,
    'Esta distinción no es semántica. Un consultor financiero analiza, diagnostica y propone. Un interim manager actúa sobre '
    'la estructura de costes, negocia con entidades financieras, cierra el reporting mensual, presenta al comité de inversión o '
    'lidera la integración del equipo adquirido. La diferencia entre recomendación y ejecución es la diferencia entre un '
    'documento y un resultado.'
)

add_body(doc,
    'El modelo de misión define esta lógica con precisión: plazos acordados, entregables concretos y transferencia ordenada '
    'del liderazgo al finalizar. El CFO interino no crea dependencia: construye capacidad y la entrega.'
)

# ════════════════════════════════════════════════════════════════════════════
# BLOQUE 5 — H2 2
# ════════════════════════════════════════════════════════════════════════════
add_h2(doc, 'Las cinco situaciones en las que el CFO interino es la respuesta correcta')

add_body(doc,
    'El 44% de las empresas españolas ha redefinido sus planes de crecimiento. Detrás de esa cifra hay decisiones que '
    'requieren liderazgo financiero ejecutivo inmediato, no en seis meses.'
)

add_body(doc,
    'La primera situación es la vacante urgente. El proceso de selección permanente de un CFO dura en España entre cuatro y '
    'nueve meses para perfiles directivos senior. Un CFO interino puede incorporarse en dos a cuatro semanas. La empresa no '
    'pierde el cierre, el reporting ni el control durante ese período.'
)

add_body(doc,
    'La segunda es la reestructuración financiera o el turnaround. Un interim manager de turnaround tiene experiencia '
    'específica en negociación de deuda, gestión de liquidez bajo presión y comunicación con accionistas e inversores en '
    'situaciones de crisis. No es un perfil genérico: es un directivo que ha estado en esa misma sala con esa misma tensión.'
)

add_body(doc,
    'La tercera es el fundraising o la ronda de financiación. Preparar un data room, construir el modelo financiero que '
    'resiste el due diligence, y representar a la empresa ante inversores requiere un perfil que hable el idioma del capital. '
    'El director financiero temporal especializado en fundraising reduce el tiempo de cierre y aumenta la calidad del proceso.'
)

add_body(doc,
    'La cuarta es la integración post-adquisición. Los 100 días posteriores a un cierre de M&A son los más críticos para la '
    'función financiera. En carteras de private equity, el interim management para estos períodos se ha convertido en práctica '
    'estándar porque la velocidad de integración determina la creación de valor.'
)

add_body(doc,
    'La quinta es el crecimiento acelerado o la internacionalización. Cuando una empresa escala rápido, la función financiera '
    'suele quedarse atrás: los sistemas, los equipos y los procesos no crecen al mismo ritmo que la facturación. Un CFO '
    'interino con experiencia en escalado puede construir la estructura financiera que el crecimiento exige sin comprometer la '
    'operación del día a día.'
)

# CTA secundario inline
p_cta2 = doc.add_paragraph()
p_cta2.style = 'Normal'
p_cta2.paragraph_format.space_before = Pt(6)
p_cta2.paragraph_format.space_after  = Pt(12)
run_cta2 = p_cta2.add_run(
    '¿Reconoces tu empresa en alguno de estos escenarios? Manager in Motion puede presentarte un CFO interino cualificado '
    'en menos de 72 horas.'
)
run_cta2.bold = True
run_cta2.font.size = Pt(11)
run_cta2.font.name = 'Arial'

# ════════════════════════════════════════════════════════════════════════════
# BLOQUE 6 — H2 3
# ════════════════════════════════════════════════════════════════════════════
add_h2(doc, 'El mercado europeo y la ecuación económica del CFO interino en España')

add_body(doc,
    'Los mercados europeos más maduros en interim management ofrecen una señal adelantada sobre lo que está llegando a España. '
    'En el Reino Unido, el 51% de todas las solicitudes de liderazgo interino corresponden a la función financiera. En Francia, '
    'el DAF de transición —equivalente al CFO interino— es el perfil más demandado dos años consecutivos. En Alemania, el '
    'mercado de interim management alcanzó un volumen récord de 2.700 millones de euros, con una tarifa media de 1.317 euros '
    'por día, y la función financiera lidera la demanda. España sigue esta curva con un desfase estimado de 12 a 18 meses: '
    'quien publica autoridad ahora, ocupa posición cuando la demanda española alcance ese nivel.'
)

add_body(doc,
    'En términos de coste, la comparación con la contratación permanente no es la que parece a primera vista. Las tarifas de '
    'referencia para un CFO interino de nivel ejecutivo en Europa se sitúan entre 1.000 y 1.500 euros por día. A diferencia de '
    'una contratación permanente, no generan coste de despido, ni cargas sociales estructurales, ni el riesgo de un proceso de '
    'selección fallido. El coste real de una vacante de CFO sin cubrir —en decisiones no tomadas, covenants incumplidos, '
    'rondas retrasadas o equipos desmotivados— supera con frecuencia el coste de la misión interina.'
)

add_body(doc,
    'El modelo también incorpora una lógica que está creciendo entre un 30% y un 35% en el mercado británico: el '
    '"try-before-you-buy". La empresa trabaja con el interim manager durante la misión y, si el encaje cultural y estratégico '
    'es el correcto, puede incorporarlo de forma permanente al final del proceso. El CFO interino se convierte así en el mejor '
    'proceso de selección que existe: con resultados reales como referencia.'
)

add_body(doc,
    'Seleccionar al perfil adecuado requiere criterios específicos: sector de actividad, tamaño y complejidad de la empresa, '
    'tipo de misión y momento del ciclo de vida del negocio. La diferencia entre una firma especializada en interim management '
    'y un headhunter generalista no es solo de red: es de metodología de diagnóstico, de velocidad de presentación de '
    'candidatos y de acompañamiento durante la misión. Las primeras 72 horas del CFO interino determinan su capacidad de '
    'impacto: diagnóstico financiero rápido, alineación con el CEO y priorización de los tres a cinco focos de actuación '
    'inmediata.'
)

# ════════════════════════════════════════════════════════════════════════════
# BLOQUE 7 — FAQ
# ════════════════════════════════════════════════════════════════════════════
add_h2(doc, 'Preguntas frecuentes sobre el CFO interino en España')

# FAQ 1
add_h3(doc, '¿Qué es exactamente un CFO interino y en qué se diferencia de un consultor financiero?')
add_body(doc,
    'El CFO interino es un directivo financiero senior que se incorpora a la empresa con autoridad ejecutiva real: ocupa el '
    'cargo de director financiero, forma parte del comité de dirección y responde ante el CEO o el consejo de administración. '
    'A diferencia de un consultor, no recomienda ni entrega informes: ejecuta, decide y es responsable de los resultados '
    'financieros durante el tiempo que dura la misión.'
)

# FAQ 2
add_h3(doc, '¿Cuánto tiempo dura una misión de CFO interino en España?')
add_body(doc,
    'La duración habitual oscila entre 3 y 18 meses, según el tipo de misión. Las coberturas de vacante urgente suelen '
    'resolverse en 3 a 6 meses. Las misiones de reestructuración, transformación financiera o preparación para una ronda de '
    'inversión pueden extenderse entre 9 y 18 meses.'
)

# FAQ 3
add_h3(doc, '¿Cuánto cuesta contratar un director financiero interino?')
add_body(doc,
    'El coste varía según el perfil del directivo, el tamaño de la empresa y la duración de la misión. En el mercado europeo, '
    'las tarifas de referencia para un CFO interino de nivel ejecutivo se sitúan entre 1.000 y 1.500 euros por día. A '
    'diferencia de una contratación permanente, no conlleva coste de despido, ni cargas sociales estructurales, ni proceso de '
    'selección de seis meses.'
)

# FAQ 4
add_h3(doc, '¿Puede un CFO interino gestionar una reestructuración de deuda o un turnaround?')
add_body(doc,
    'Sí. Es uno de los casos de uso más exigentes y más frecuentes del interim management financiero. El CFO interino de '
    'turnaround tiene experiencia específica en negociación con entidades financieras, gestión de liquidez bajo presión y '
    'comunicación con accionistas e inversores en situaciones de crisis. No es un perfil genérico: debe seleccionarse con '
    'criterios muy precisos según el tipo de situación.'
)

# FAQ 5
add_h3(doc, '¿Cuándo tiene más sentido un CFO interino que una contratación permanente?')
add_body(doc,
    'El CFO interino es la opción más racional cuando la empresa necesita capacidad financiera ejecutiva de alto nivel de forma '
    'inmediata y no puede comprometerse con una contratación permanente a largo plazo. Las cinco situaciones más frecuentes: '
    'vacante urgente, reestructuración, fundraising, integración post-M&A y crecimiento acelerado sin equipo financiero que '
    'escale.'
)

# FAQ 6
add_h3(doc, '¿Cómo de rápido puede incorporarse un CFO interino en una empresa española?')
add_body(doc,
    'Un CFO interino puede incorporarse en un plazo de 2 a 4 semanas desde el primer contacto con la firma especializada, en '
    'condiciones normales. En situaciones de urgencia máxima, algunas firmas pueden presentar el perfil adecuado en menos de '
    '72 horas. Esta velocidad es estructuralmente imposible en un proceso de selección permanente, que en España tiene una '
    'duración media de 4 a 9 meses para perfiles directivos senior.'
)

# ════════════════════════════════════════════════════════════════════════════
# BLOQUE 8 — CTA PRINCIPAL
# ════════════════════════════════════════════════════════════════════════════
add_separator(doc)

p_cta_title = doc.add_paragraph()
p_cta_title.style = 'Normal'
run_ct = p_cta_title.add_run('¿Tu empresa necesita un director financiero interino?')
run_ct.bold = True
run_ct.font.size = Pt(13)
run_ct.font.name = 'Arial'
run_ct.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)
p_cta_title.paragraph_format.space_after = Pt(6)

add_body(doc,
    'En Manager in Motion evaluamos tu situación y te presentamos el perfil de CFO interino más adecuado para tu contexto. '
    'La primera conversación no tiene coste ni compromiso.'
)

p_btn = doc.add_paragraph()
p_btn.style = 'Normal'
run_btn = p_btn.add_run('→ Solicitar consulta sin compromiso')
run_btn.bold = True
run_btn.font.size = Pt(12)
run_btn.font.name = 'Arial'
run_btn.font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)
p_btn.paragraph_format.space_after = Pt(16)

# ════════════════════════════════════════════════════════════════════════════
# BLOQUE 9 — SUGERENCIAS DE ENLAZADO INTERNO
# ════════════════════════════════════════════════════════════════════════════
add_separator(doc)

p_enc_title = doc.add_paragraph()
run_enc = p_enc_title.add_run('Sugerencias de enlazado interno')
run_enc.bold = True
run_enc.font.size = Pt(11)
run_enc.font.name = 'Arial'
run_enc.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
p_enc_title.paragraph_format.space_after = Pt(4)

enlaces_internos = [
    '1. Página de servicio: Dirección Financiera Interina / CFO Interino — ancla sugerida: "directivos financieros interinos especializados" o "servicio de CFO interino de Manager in Motion". Colocar en H2 "Las cinco situaciones" y en el CTA.',
    '2. Página de servicio: Interim Management para Private Equity — ancla: "interim management para fondos de private equity". Colocar en el párrafo de integración post-adquisición.',
    '3. Página "¿Qué es el interim management?" — ancla: "interim management". Colocar en la introducción o en el H2 de definición.',
    '4. Página de servicio: Dirección General Interina (CEO Interino) — ancla: "director general interino" o "interim CEO". Colocar donde se mencione la relación CFO-CEO.',
    '5. Página de contacto / formulario de consulta — ancla: "habla con un especialista en dirección financiera interina". Colocar en la conclusión y en el CTA.',
]
for enlace in enlaces_internos:
    p_e = doc.add_paragraph()
    p_e.style = 'Normal'
    run_e = p_e.add_run(enlace)
    run_e.font.size = Pt(10)
    run_e.font.name = 'Arial'
    run_e.font.color.rgb = RGBColor(0x60, 0x60, 0x60)
    p_e.paragraph_format.space_after = Pt(4)

# ════════════════════════════════════════════════════════════════════════════
# BLOQUE 10 — SEPARADOR EDITORIAL + FUENTES
# ════════════════════════════════════════════════════════════════════════════
add_separator(doc)

p_editorial = doc.add_paragraph()
run_ed = p_editorial.add_run('FUERA DEL ARTÍCULO — Solo para uso editorial')
run_ed.bold = True
run_ed.italic = True
run_ed.font.size = Pt(10)
run_ed.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
run_ed.font.name = 'Arial'
p_editorial.paragraph_format.space_after = Pt(8)

add_h2(doc, 'Fuentes consultadas')

fuentes = [
    '1. Heidrick & Struggles — Talent Lens 2026 (heidrick.com). Datos utilizados: +103% de crecimiento interanual en demanda de CFOs interinos en Europa; 51% de todas las solicitudes de liderazgo interino en UK corresponden a la función financiera; crecimiento del modelo "try-before-you-buy" entre 30% y 35% en 2026 en UK.',
    '2. DDIM — Marktstudie 2026 (ddim.de). Datos utilizados: mercado alemán de interim management valorado en 2.700 millones de euros; tarifa media de 1.317 euros/día.',
    '3. PwC España — Encuesta de directivos 2025/2026 (pwc.es). Datos utilizados: 98% de directivos españoles señalan la escasez de talento directivo como el principal freno a la productividad.',
    '4. KPMG España — Perspectivas 2026 (kpmg.com/es). Datos utilizados: 44% de empresas españolas ha redefinido sus planes de crecimiento.',
    '5. EY España — Perspectivas del mercado de trabajo 2026 (ey.com/es_ES). Contexto utilizado: vulnerabilidad de las pymes españolas ante incrementos de costes laborales como argumento complementario al modelo interino.',
]
for fuente in fuentes:
    p_f = doc.add_paragraph()
    p_f.style = 'Normal'
    run_f = p_f.add_run(fuente)
    run_f.font.size = Pt(10)
    run_f.font.name = 'Arial'
    p_f.paragraph_format.space_after = Pt(4)

doc.add_paragraph()

p_notas_title = doc.add_paragraph()
run_nt = p_notas_title.add_run('Notas de revisión final')
run_nt.bold = True
run_nt.font.size = Pt(11)
run_nt.font.name = 'Arial'
p_notas_title.paragraph_format.space_after = Pt(4)

notas = [
    '— El dato del 103% (Heidrick & Struggles, Talent Lens 2026) debe verificarse con la publicación original antes de publicar; la fecha de acceso al informe puede requerir actualización.',
    '— El dato del 98% de PwC España proviene del informe de directivos más reciente disponible; verificar que la edición 2025/2026 mantiene este porcentaje.',
    '— El rango de tarifas (1.000–1.500 €/día) refleja el mercado europeo de referencia; ajustar si Manager in Motion dispone de datos específicos del mercado español.',
    '— Confirmar las URLs de enlazado interno (páginas de servicio, formulario de contacto) con el equipo web antes de publicar.',
    '— El esquema FAQ (FAQPage schema) debe implementarse en el CMS antes de la publicación para optar a rich snippets en Google.',
]
for nota in notas:
    p_n = doc.add_paragraph()
    p_n.style = 'Normal'
    run_n = p_n.add_run(nota)
    run_n.font.size = Pt(10)
    run_n.font.name = 'Arial'
    run_n.font.color.rgb = RGBColor(0x60, 0x60, 0x60)
    p_n.paragraph_format.space_after = Pt(4)

# ════════════════════════════════════════════════════════════════════════════
# GUARDAR
# ════════════════════════════════════════════════════════════════════════════
output_path = '/home/user/mimo-seo/output/agent3_article.docx'
doc.save(output_path)
print(f'Archivo generado: {output_path}')

import os
size = os.path.getsize(output_path)
print(f'Tamaño del archivo: {size:,} bytes ({size/1024:.1f} KB)')
